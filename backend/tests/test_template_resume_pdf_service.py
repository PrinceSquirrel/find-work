from __future__ import annotations

from io import BytesIO
from zipfile import ZipFile

from docx import Document
from PIL import Image
from pypdf import PdfReader
from reportlab.pdfgen import canvas

from app.services import pdf_service
from app.services.pdf_service import DocxProjectTemplateService, PdfRenderValidator, TailoredResumePdfService


def _docx_with_editable_body_and_image() -> bytes:
    image_buffer = BytesIO()
    Image.new("RGB", (8, 8), color="white").save(image_buffer, format="PNG")
    image_buffer.seek(0)
    document = Document()
    document.add_paragraph("胡俊")
    document.add_paragraph("电话: 15822716099")
    document.add_paragraph("教育经历: 吉林大学")
    document.add_picture(image_buffer)
    document.add_paragraph("技能: Python, SQL")
    document.add_paragraph("项目经历")
    document.add_paragraph("旧项目描述，应该被替换。")
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def _paragraphs(docx_bytes: bytes) -> list[str]:
    return [paragraph.text for paragraph in Document(BytesIO(docx_bytes)).paragraphs]


def _media_names(docx_bytes: bytes) -> list[str]:
    with ZipFile(BytesIO(docx_bytes)) as archive:
        return sorted(name for name in archive.namelist() if name.startswith("word/media/"))


def _pdf_with_pages(page_count: int) -> bytes:
    buffer = BytesIO()
    pdf = canvas.Canvas(buffer)
    for index in range(page_count):
        pdf.drawString(72, 720, f"page {index + 1}")
        pdf.showPage()
    pdf.save()
    return buffer.getvalue()


def test_docx_template_preserves_identity_and_education_then_replaces_editable_body():
    template_bytes = _docx_with_editable_body_and_image()

    result = DocxProjectTemplateService().render(
        template_bytes,
        "技能: Python, SQL, 数据分析\n项目经历: 围绕岗位突出 Python、SQL 和数据看板。",
    )

    paragraphs = _paragraphs(result)
    assert "胡俊" in paragraphs
    assert "电话: 15822716099" in paragraphs
    assert "教育经历: 吉林大学" in paragraphs
    assert "技能: Python, SQL" not in paragraphs
    assert "旧项目描述，应该被替换。" not in paragraphs
    assert "技能: Python, SQL, 数据分析\n项目经历: 围绕岗位突出 Python、SQL 和数据看板。" in paragraphs
    assert _media_names(result) == _media_names(template_bytes)


def test_tailored_resume_pdf_service_retries_with_compressed_project_until_one_page():
    class FakeConverter:
        def __init__(self):
            self.docx_inputs: list[bytes] = []

        def convert(self, docx_bytes: bytes) -> bytes:
            self.docx_inputs.append(docx_bytes)
            return _pdf_with_pages(2 if len(self.docx_inputs) == 1 else 1)

    converter = FakeConverter()
    service = TailoredResumePdfService(converter=converter)
    bundle = {
        "id": 1,
        "resume_rewrite": "\n".join(f"第 {index} 行简历正文，需要被压缩。" for index in range(20)),
        "project_rewrite": "旧项目改写字段不应优先使用。",
    }

    pdf_bytes = service.render(bundle, _docx_with_editable_body_and_image())

    assert len(converter.docx_inputs) == 2
    assert len(PdfReader(BytesIO(pdf_bytes)).pages) == 1


def test_pdf_validator_rejects_invalid_pdf_bytes_with_clear_error():
    validator = PdfRenderValidator()

    try:
        validator.validate(b"not a pdf", max_pages=1)
    except RuntimeError as exc:
        message = str(exc)
    else:
        raise AssertionError("Expected invalid PDF bytes to fail validation.")

    assert "PDF" in message
    assert "invalid" in message.lower()


def test_pdf_validator_skips_default_renderer_when_wrapper_is_unavailable(monkeypatch):
    class FailedProbe:
        returncode = 1
        stdout = ""
        stderr = "The system cannot find the path specified."

    monkeypatch.setattr(pdf_service.shutil, "which", lambda name: "pdftoppm.cmd")
    monkeypatch.setattr(pdf_service.subprocess, "run", lambda *args, **kwargs: FailedProbe())

    assert PdfRenderValidator().validate(_pdf_with_pages(1), max_pages=1) == 1


def test_tailored_resume_pdf_service_validates_render_before_returning_pdf():
    class FakeConverter:
        def convert(self, docx_bytes: bytes) -> bytes:
            return _pdf_with_pages(1)

    class RecordingValidator:
        def __init__(self):
            self.calls: list[tuple[bytes, int]] = []

        def validate(self, pdf_bytes: bytes, max_pages: int) -> int:
            self.calls.append((pdf_bytes, max_pages))
            return 1

    validator = RecordingValidator()
    service = TailoredResumePdfService(converter=FakeConverter(), validator=validator)

    pdf_bytes = service.render(
        {"resume_rewrite": "鎶€鑳? Python\n缁忓巻: Agent 宸ヤ綔鍙?"},
        _docx_with_editable_body_and_image(),
    )

    assert pdf_bytes.startswith(b"%PDF")
    assert len(validator.calls) == 1
    assert validator.calls[0][1] == 1


def test_docx_template_missing_editable_body_error_is_not_project_only():
    document = Document()
    document.add_paragraph("胡俊")
    document.add_paragraph("电话: 15822716099")
    output = BytesIO()
    document.save(output)

    try:
        DocxProjectTemplateService().render(output.getvalue(), "技能: Python\n经历: 数据分析")
    except ValueError as exc:
        message = str(exc)
    else:
        raise AssertionError("Expected missing editable body to fail clearly.")

    assert "可编辑简历正文" in message
    assert "项目经历" not in message
