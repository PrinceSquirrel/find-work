import base64
import inspect
import json
import sqlite3
from io import BytesIO

import pytest
from fastapi.testclient import TestClient

from app.main import create_app
from app.schemas import ExtractedJobCandidate, PlatformJobExtraction, PlatformSession, PlatformSessionsResponse
from app.services import browser_job_extractor_service
from app.services import job_application_service
from app.services import platform_session_service


@pytest.fixture(autouse=True)
def _isolate_external_env_file(tmp_path, monkeypatch):
    monkeypatch.setenv("AGENT_BUSINESS_ENV_FILE", str(tmp_path / "missing-agent-business.env"))


def _one_page_pdf(text: str = "template pdf") -> bytes:
    from reportlab.pdfgen import canvas

    buffer = BytesIO()
    pdf = canvas.Canvas(buffer)
    for index, line in enumerate(text.splitlines() or ["template pdf"]):
        pdf.drawString(72, 720 - index * 18, line[:110])
    pdf.showPage()
    pdf.save()
    return buffer.getvalue()


def _one_pixel_png() -> bytes:
    return base64.b64decode(
        "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMCAO+/p9sAAAAASUVORK5CYII="
    )


def _create_demo_application(client: TestClient) -> int:
    resume_text = (
        "赵六\n"
        "技能: Python, FastAPI, React, SQL, 数据分析\n"
        "项目: 求职 Agent 工作台，负责后端 API 和前端看板。"
    )
    upload_response = client.post(
        "/api/resumes",
        files={"file": ("resume.txt", resume_text.encode("utf-8"), "text/plain")},
    )
    resume_id = upload_response.json()["id"]

    search_response = client.post(
        "/api/search-runs",
        json={
            "resume_id": resume_id,
            "keywords": ["Python 实习"],
            "city": "上海",
            "platforms": ["boss"],
        },
    )
    assert search_response.status_code == 201

    jobs_response = client.get("/api/jobs")
    job_id = jobs_response.json()[0]["id"]
    record = client.app.state.service.create_application(job_id, note="测试夹具：平台已确认投递")
    return record.id


def test_full_resume_to_application_flow(tmp_path, monkeypatch):
    class FakeBrowserJobExtractorService:
        def apply_to_job(self, platform, url):
            return {
                "confirmed": True,
                "status": "applied",
                "action": "clicked_apply",
                "evidence": f"{platform} confirmed the application",
                "source_url": url,
            }

    monkeypatch.setattr(job_application_service, "BrowserJobExtractorService", FakeBrowserJobExtractorService)
    app = create_app(db_path=tmp_path / "agent-business.sqlite3")
    client = TestClient(app)

    resume_text = (
        "李四\n"
        "教育经历: 浙江大学 计算机科学\n"
        "技能: Python, FastAPI, React, SQL, 数据分析\n"
        "项目: 多 Agent 求职助手，负责后端 API、前端看板和模型调用统计。"
    )

    upload_response = client.post(
        "/api/resumes",
        files={"file": ("resume.txt", resume_text.encode("utf-8"), "text/plain")},
    )
    assert upload_response.status_code == 201
    resume_id = upload_response.json()["id"]

    search_response = client.post(
        "/api/search-runs",
        json={
            "resume_id": resume_id,
            "keywords": ["Python 实习", "Agent 实习"],
            "city": "上海",
            "platforms": ["boss", "shixiseng"],
        },
    )
    assert search_response.status_code == 201
    assert search_response.json()["status"] == "completed"

    jobs_response = client.get("/api/jobs")
    assert jobs_response.status_code == 200
    jobs = jobs_response.json()
    assert len(jobs) >= 2
    assert {job["platform"] for job in jobs} == {"boss", "shixiseng"}
    assert all("match" in job for job in jobs)

    target_job_id = jobs[0]["id"]
    tailor_response = client.post(f"/api/jobs/{target_job_id}/tailor", json={"resume_id": resume_id})
    assert tailor_response.status_code == 201
    tailored = tailor_response.json()
    assert tailored["truth_check_passed"] is True
    assert tailored["greeting"]["message"]

    apply_response = client.post(
        f"/api/jobs/{target_job_id}/platform-apply",
        json={"note": "用户确认后投递"},
    )
    assert apply_response.status_code == 201
    application_id = apply_response.json()["id"]

    status_response = client.patch(
        f"/api/applications/{application_id}/status",
        json={"status": "read", "note": "半自动同步：招聘方已读"},
    )
    assert status_response.status_code == 200
    assert status_response.json()["current_status"] == "read"

    analytics_response = client.get("/api/analytics/applications")
    assert analytics_response.status_code == 200
    analytics = analytics_response.json()
    assert analytics["totals"]["applications"] == 1
    assert analytics["totals"]["read_rate"] == 1.0
    assert "hourly" in analytics

    metrics_response = client.get("/api/metrics/llm-usage")
    assert metrics_response.status_code == 200
    assert metrics_response.json()["total_tokens"] > 0


def test_pdf_resume_upload_can_generate_tailored_materials(tmp_path):
    app = create_app(db_path=tmp_path / "agent-business.sqlite3")
    client = TestClient(app)

    pdf_bytes = _one_page_pdf(
        "Skills: Python, FastAPI, React, SQL, Agent, API\n"
        "Project: local job application agent dashboard and model cost metrics"
    )
    upload_response = client.post(
        "/api/resumes",
        files={"file": ("resume.pdf", pdf_bytes, "application/pdf")},
    )

    assert upload_response.status_code == 201
    resume_payload = upload_response.json()
    assert resume_payload["file_type"] == "pdf"
    assert resume_payload["template_available"] is False
    assert resume_payload["profile"]["pdf_reading"]["method"] == "pypdf"
    assert resume_payload["profile"]["pdf_reading"]["status"] == "success"
    assert resume_payload["profile"]["can_generate_materials"] is True

    search_response = client.post(
        "/api/search-runs",
        json={
            "resume_id": resume_payload["id"],
            "keywords": ["Python internship"],
            "city": "Shanghai",
            "platforms": ["boss"],
        },
    )
    assert search_response.status_code == 201
    job_id = client.get("/api/jobs").json()[0]["id"]

    tailor_response = client.post(f"/api/jobs/{job_id}/tailor", json={"resume_id": resume_payload["id"]})

    assert tailor_response.status_code == 201
    tailored = tailor_response.json()
    assert tailored["resume_id"] == resume_payload["id"]
    assert tailored["greeting"]["message"]


def test_scanned_pdf_upload_requires_manual_text_before_materials(tmp_path):
    app = create_app(db_path=tmp_path / "agent-business.sqlite3")
    client = TestClient(app)

    upload_response = client.post(
        "/api/resumes",
        files={"file": ("resume.pdf", _one_page_pdf(" "), "application/pdf")},
    )

    assert upload_response.status_code == 201
    payload = upload_response.json()
    extraction = payload["profile"]["pdf_reading"]
    assert extraction["source_type"] == "pdf_scan"
    assert extraction["status"] == "needs_ocr"
    assert extraction["manual_text_required"] is True
    assert payload["profile"]["can_generate_materials"] is False

    manual_response = client.patch(
        f"/api/resumes/{payload['id']}/manual-text",
        json={"raw_text": "技能: Python, FastAPI, React\n城市: 上海\n项目: 求职 Agent 工作台"},
    )

    assert manual_response.status_code == 200
    manual_payload = manual_response.json()
    assert "Python" in manual_payload["raw_text"]
    assert manual_payload["profile"]["extraction"]["source_type"] == "manual"
    assert manual_payload["profile"]["extraction"]["manual_text_required"] is False
    assert manual_payload["profile"]["can_generate_materials"] is True
    assert manual_payload["profile"]["suggested_city"] == "上海"


def test_image_resume_upload_keeps_image_and_requests_manual_text(tmp_path):
    db_path = tmp_path / "agent-business.sqlite3"
    app = create_app(db_path=db_path)
    client = TestClient(app)
    image_bytes = _one_pixel_png()

    response = client.post(
        "/api/resumes",
        files={"file": ("resume.png", image_bytes, "image/png")},
    )

    assert response.status_code == 201
    payload = response.json()
    assert payload["file_type"] == "png"
    assert payload["profile"]["image_reading"]["manual_text_required"] is True
    assert payload["profile"]["can_generate_materials"] is False
    with sqlite3.connect(db_path) as conn:
        row = conn.execute(
            "SELECT file_type, original_file_bytes FROM resumes WHERE id = ?",
            (payload["id"],),
        ).fetchone()
    assert row[0] == "png"
    assert row[1] == image_bytes


def test_docx_resume_upload_preserves_original_template_bytes(tmp_path):
    from docx import Document

    document = Document()
    document.add_paragraph("张三")
    document.add_paragraph("项目经历")
    document.add_paragraph("求职 Agent 工作台，负责 Python 和 React。")
    buffer = BytesIO()
    document.save(buffer)
    content = buffer.getvalue()
    db_path = tmp_path / "agent-business.sqlite3"
    app = create_app(db_path=db_path)
    client = TestClient(app)

    response = client.post(
        "/api/resumes",
        files={"file": ("resume.docx", content, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
    )

    assert response.status_code == 201
    payload = response.json()
    assert payload["file_type"] == "docx"
    assert payload["template_available"] is True
    with sqlite3.connect(db_path) as conn:
        row = conn.execute(
            "SELECT file_type, template_available, original_file_bytes FROM resumes WHERE id = ?",
            (payload["id"],),
        ).fetchone()
    assert row[0] == "docx"
    assert row[1] == 1
    assert row[2] == content


def test_platform_apply_creates_application_only_after_platform_confirmation(tmp_path, monkeypatch):
    class FakeBrowserJobExtractorService:
        calls = []

        def apply_to_job(self, platform, url):
            self.calls.append({"platform": platform, "url": url})
            return {
                "confirmed": True,
                "status": "applied",
                "action": "clicked_apply",
                "evidence": "clicked platform button: 立即沟通",
                "button_text": "立即沟通",
                "page_summary": "平台页面显示已沟通",
                "source_url": url,
            }

    monkeypatch.setattr(job_application_service, "BrowserJobExtractorService", FakeBrowserJobExtractorService)
    app = create_app(db_path=tmp_path / "agent-business.sqlite3")
    client = TestClient(app)
    resume_response = client.post(
        "/api/resumes",
        files={"file": ("resume.txt", "技能: Python, FastAPI".encode("utf-8"), "text/plain")},
    )
    run_response = client.post(
        "/api/search-runs",
        json={
            "resume_id": resume_response.json()["id"],
            "keywords": ["Python 实习"],
            "city": "上海",
            "platforms": ["boss"],
        },
    )
    job = client.get(f"/api/jobs?search_run_id={run_response.json()['id']}").json()[0]

    response = client.post(f"/api/jobs/{job['id']}/platform-apply", json={"note": "用户确认在 BOSS 投递"})

    assert response.status_code == 201
    payload = response.json()
    assert payload["current_status"] == "applied"
    assert "clicked platform button" in payload["latest_note"]
    assert payload["platform_proof"]["platform"] == "boss"
    assert payload["platform_proof"]["source_url"] == job["url"]
    assert payload["platform_proof"]["button_text"] == "立即沟通"
    assert payload["platform_proof"]["status"] == "applied"
    assert payload["platform_proof"]["action"] == "clicked_apply"
    assert payload["platform_proof"]["evidence"] == "clicked platform button: 立即沟通"
    assert payload["platform_proof"]["page_summary"] == "平台页面显示已沟通"
    assert payload["platform_proof"]["confirmed_at"]
    assert FakeBrowserJobExtractorService.calls == [{"platform": "boss", "url": job["url"]}]
    applications = client.get("/api/applications").json()
    assert len(applications) == 1
    assert applications[0]["id"] == payload["id"]
    assert applications[0]["platform_proof"] == payload["platform_proof"]


def test_platform_apply_does_not_create_application_without_platform_confirmation(tmp_path, monkeypatch):
    class FakeBrowserJobExtractorService:
        def apply_to_job(self, platform, url):
            return {
                "confirmed": False,
                "status": "button_not_found",
                "action": "none",
                "evidence": "no platform apply or chat button was found",
                "source_url": url,
            }

    monkeypatch.setattr(job_application_service, "BrowserJobExtractorService", FakeBrowserJobExtractorService)
    app = create_app(db_path=tmp_path / "agent-business.sqlite3")
    client = TestClient(app)
    resume_response = client.post(
        "/api/resumes",
        files={"file": ("resume.txt", "技能: Python, FastAPI".encode("utf-8"), "text/plain")},
    )
    run_response = client.post(
        "/api/search-runs",
        json={
            "resume_id": resume_response.json()["id"],
            "keywords": ["Python 实习"],
            "city": "上海",
            "platforms": ["boss"],
        },
    )
    job = client.get(f"/api/jobs?search_run_id={run_response.json()['id']}").json()[0]

    response = client.post(f"/api/jobs/{job['id']}/platform-apply", json={"note": "用户确认在 BOSS 投递"})

    assert response.status_code == 409
    assert "no platform apply" in response.json()["detail"]
    assert client.get("/api/applications").json() == []


def test_platform_apply_preview_reports_button_without_creating_application(tmp_path, monkeypatch):
    class FakeBrowserJobExtractorService:
        calls = []

        def preview_apply_to_job(self, platform, url):
            self.calls.append({"platform": platform, "url": url})
            return {
                "ready": True,
                "status": "ready",
                "action": "preview",
                "button_text": "立即沟通",
                "evidence": "found platform button: 立即沟通",
                "source_url": url,
            }

    monkeypatch.setattr(job_application_service, "BrowserJobExtractorService", FakeBrowserJobExtractorService)
    app = create_app(db_path=tmp_path / "agent-business.sqlite3")
    client = TestClient(app)
    resume_response = client.post(
        "/api/resumes",
        files={"file": ("resume.txt", "技能: Python, FastAPI".encode("utf-8"), "text/plain")},
    )
    run_response = client.post(
        "/api/search-runs",
        json={
            "resume_id": resume_response.json()["id"],
            "keywords": ["Python 实习"],
            "city": "上海",
            "platforms": ["boss"],
        },
    )
    job = client.get(f"/api/jobs?search_run_id={run_response.json()['id']}").json()[0]

    response = client.post(f"/api/jobs/{job['id']}/platform-apply-preview")

    assert response.status_code == 200
    payload = response.json()
    assert payload["ready"] is True
    assert payload["button_text"] == "立即沟通"
    assert payload["job"]["id"] == job["id"]
    assert FakeBrowserJobExtractorService.calls == [{"platform": "boss", "url": job["url"]}]
    assert client.get("/api/applications").json() == []


def test_legacy_apply_record_endpoint_rejects_local_only_application_records(tmp_path):
    app = create_app(db_path=tmp_path / "agent-business.sqlite3")
    client = TestClient(app)
    resume_response = client.post(
        "/api/resumes",
        files={"file": ("resume.txt", "技能: Python, FastAPI".encode("utf-8"), "text/plain")},
    )
    run_response = client.post(
        "/api/search-runs",
        json={
            "resume_id": resume_response.json()["id"],
            "keywords": ["Python 实习"],
            "city": "上海",
            "platforms": ["boss"],
        },
    )
    job = client.get(f"/api/jobs?search_run_id={run_response.json()['id']}").json()[0]

    response = client.post(f"/api/jobs/{job['id']}/apply-record", json={"note": "只在本地记录"})

    assert response.status_code == 409
    assert "platform-apply" in response.json()["detail"]
    assert client.get("/api/applications").json() == []


def test_latest_resume_endpoint_recovers_after_app_restart(tmp_path):
    db_path = tmp_path / "agent-business.sqlite3"
    first_client = TestClient(create_app(db_path=db_path))
    first_client.post(
        "/api/resumes",
        files={"file": ("old.txt", "技能: SQL".encode("utf-8"), "text/plain")},
    )
    latest_upload = first_client.post(
        "/api/resumes",
        files={"file": ("latest.txt", "技能: Python, FastAPI".encode("utf-8"), "text/plain")},
    )
    latest_id = latest_upload.json()["id"]

    restarted_client = TestClient(create_app(db_path=db_path))
    response = restarted_client.get("/api/resumes/latest")

    assert response.status_code == 200
    payload = response.json()
    assert payload["id"] == latest_id
    assert payload["filename"] == "latest.txt"
    assert "Python" in payload["raw_text"]


def test_knowledge_reindex_and_rag_query_return_source_chunks(tmp_path):
    app = create_app(db_path=tmp_path / "agent-business.sqlite3")
    client = TestClient(app)

    upload_response = client.post(
        "/api/resumes",
        files={
            "file": (
                "resume.txt",
                "技能: Python, FastAPI, React, SQL\n项目: Agent 工作台，负责 RAG 检索和投递看板。".encode("utf-8"),
                "text/plain",
            )
        },
    )
    resume_id = upload_response.json()["id"]
    run_response = client.post(
        "/api/search-runs",
        json={
            "resume_id": resume_id,
            "keywords": ["Python 实习"],
            "city": "上海",
            "platforms": ["boss"],
        },
    )
    assert run_response.status_code == 201

    reindex_response = client.post("/api/knowledge/reindex")

    assert reindex_response.status_code == 200
    reindex_payload = reindex_response.json()
    assert reindex_payload["status"] == "completed"
    assert reindex_payload["documents"] >= 2
    assert reindex_payload["chunks"] >= 2

    documents_response = client.get("/api/knowledge/documents")
    assert documents_response.status_code == 200
    documents = documents_response.json()
    assert {"resume", "job"} <= {document["source_type"] for document in documents}

    query_response = client.post("/api/rag/query", json={"query": "FastAPI SQL Agent", "limit": 5})

    assert query_response.status_code == 200
    payload = query_response.json()
    assert payload["hits"]
    assert all(hit["source_type"] in {"resume", "job"} for hit in payload["hits"])
    assert all(hit["source_id"] for hit in payload["hits"])
    assert "基于本地知识库" in payload["answer"]


def test_skill_registry_lists_allowlist_and_enforces_high_risk_confirmation(tmp_path):
    app = create_app(db_path=tmp_path / "agent-business.sqlite3")
    client = TestClient(app)

    skills_response = client.get("/api/skills")

    assert skills_response.status_code == 200
    skills = {skill["id"]: skill for skill in skills_response.json()}
    assert "job.search.browser_cdp" in skills
    assert "application.apply.platform" in skills
    assert skills["application.apply.platform"]["risk_level"] == "high"
    assert skills["application.apply.platform"]["requires_confirmation"] is True

    high_risk_response = client.post(
        "/api/skills/application.apply.platform/run",
        json={"arguments": {"job_id": 123}, "confirmed": False},
    )

    assert high_risk_response.status_code == 200
    high_risk_payload = high_risk_response.json()
    assert high_risk_payload["status"] == "requires_confirmation"
    assert high_risk_payload["requires_confirmation"] is True

    unknown_response = client.post("/api/skills/not.registered/run", json={})
    assert unknown_response.status_code == 404


def test_low_risk_knowledge_reindex_skill_runs_directly(tmp_path):
    app = create_app(db_path=tmp_path / "agent-business.sqlite3")
    client = TestClient(app)
    client.post(
        "/api/resumes",
        files={"file": ("resume.txt", "技能: Python, RAG, Agent".encode("utf-8"), "text/plain")},
    )

    response = client.post("/api/skills/knowledge.reindex/run", json={"arguments": {}})

    assert response.status_code == 200
    payload = response.json()
    assert payload["status"] == "success"
    assert payload["requires_confirmation"] is False
    assert payload["result"]["documents"] >= 1
    assert payload["result"]["chunks"] >= 1


def test_mcp_gateway_defaults_to_empty_allowlist_and_rejects_unknown_tools(tmp_path, monkeypatch):
    monkeypatch.delenv("AGENT_BUSINESS_MCP_SERVERS", raising=False)
    app = create_app(db_path=tmp_path / "agent-business.sqlite3")
    client = TestClient(app)

    assert client.get("/api/mcp/servers").json() == []
    assert client.get("/api/mcp/tools").json() == []

    response = client.post("/api/mcp/tools/not-allowed/echo/call", json={"arguments": {}})

    assert response.status_code == 403
    assert "allowlist" in response.json()["detail"]


def test_mcp_gateway_lists_configured_allowlist_and_requires_confirmation(tmp_path, monkeypatch):
    monkeypatch.setenv(
        "AGENT_BUSINESS_MCP_SERVERS",
        json.dumps(
            [
                {
                    "id": "fake",
                    "command": "python",
                    "args": ["fake_mcp_server.py"],
                    "enabled": True,
                    "allowed_tools": ["echo"],
                }
            ]
        ),
    )
    app = create_app(db_path=tmp_path / "agent-business.sqlite3")
    client = TestClient(app)

    servers = client.get("/api/mcp/servers").json()
    tools = client.get("/api/mcp/tools").json()
    response = client.post("/api/mcp/tools/fake/echo/call", json={"arguments": {"text": "hi"}})

    assert servers[0]["id"] == "fake"
    assert tools[0]["server_id"] == "fake"
    assert tools[0]["name"] == "echo"
    assert response.status_code == 200
    assert response.json()["status"] == "requires_confirmation"


def test_status_endpoint_rejects_illegal_application_transition(tmp_path):
    app = create_app(db_path=tmp_path / "agent-business.sqlite3")
    client = TestClient(app)
    application_id = _create_demo_application(client)

    status_response = client.patch(
        f"/api/applications/{application_id}/status",
        json={"status": "replied", "note": "不能跳过已读直接回复"},
    )

    assert status_response.status_code == 409
    assert "Cannot transition application" in status_response.json()["detail"]

    applications_response = client.get("/api/applications")
    assert applications_response.status_code == 200
    assert applications_response.json()[0]["current_status"] == "applied"


def test_application_sync_returns_read_only_proposals_without_overwriting_status(tmp_path, monkeypatch):
    class FakeCdpResponse:
        def __enter__(self):
            return self

        def __exit__(self, exc_type, exc, traceback):
            return False

        def read(self):
            return json.dumps(
                [
                    {
                        "type": "page",
                        "url": "https://www.zhipin.com/web/geek/chat?securityId=hidden",
                        "webSocketDebuggerUrl": "ws://127.0.0.1:9222/devtools/page/1",
                    }
                ]
            ).encode("utf-8")

    def fake_evaluate(self, websocket_url, expression):
        assert websocket_url == "ws://127.0.0.1:9222/devtools/page/1"
        assert "application-sync-readonly" in expression
        return {
            "items": [
                {
                    "text": "星河智能科技 Python 实习 招聘方已回复：方便约聊吗",
                    "url": "https://www.zhipin.com/web/geek/chat?securityId=hidden",
                }
            ],
            "diagnostics": {
                "candidate_item_count": 1,
                "matched_status_keywords": {"已回复": 1},
            },
        }

    monkeypatch.delenv("BROWSER_CDP_URL", raising=False)
    monkeypatch.setattr(browser_job_extractor_service, "urlopen", lambda url, timeout: FakeCdpResponse())
    monkeypatch.setattr(browser_job_extractor_service.CdpRuntimeClient, "evaluate", fake_evaluate)
    app = create_app(db_path=tmp_path / "agent-business.sqlite3")
    client = TestClient(app)
    application_id = _create_demo_application(client)

    response = client.post("/api/applications/sync", json={"platforms": ["boss"], "limit": 20})

    assert response.status_code == 200
    payload = response.json()
    assert payload["status"] == "completed"
    assert payload["mode"] == "browser_cdp_readonly"
    assert payload["updated"] == 0
    assert payload["proposals"][0]["application_id"] == application_id
    assert payload["proposals"][0]["detected_status"] == "replied"
    assert payload["proposals"][0]["suggested_status"] == "read"
    assert payload["proposals"][0]["requires_manual_confirmation"] is True
    assert "hidden" not in str(payload)

    applications_response = client.get("/api/applications")
    assert applications_response.json()[0]["current_status"] == "applied"


def test_application_sync_reports_missing_cdp_without_updating_status(tmp_path, monkeypatch):
    monkeypatch.delenv("BROWSER_CDP_URL", raising=False)
    monkeypatch.setattr(
        browser_job_extractor_service,
        "urlopen",
        lambda url, timeout: (_ for _ in ()).throw(OSError("no local CDP")),
    )
    app = create_app(db_path=tmp_path / "agent-business.sqlite3")
    client = TestClient(app)
    application_id = _create_demo_application(client)

    response = client.post("/api/applications/sync", json={"platforms": ["boss"]})

    assert response.status_code == 200
    payload = response.json()
    assert payload["status"] == "not_configured"
    assert payload["updated"] == 0
    assert payload["proposals"] == []
    assert payload["diagnostics"][0]["failure_reason"] == "未配置 BROWSER_CDP_URL"
    assert client.get("/api/applications").json()[0]["id"] == application_id
    assert client.get("/api/applications").json()[0]["current_status"] == "applied"


def test_search_run_rejects_unknown_platforms(tmp_path):
    app = create_app(db_path=tmp_path / "agent-business.sqlite3")
    client = TestClient(app)

    upload_response = client.post(
        "/api/resumes",
        files={"file": ("resume.txt", "技能: Python, FastAPI".encode("utf-8"), "text/plain")},
    )
    resume_id = upload_response.json()["id"]

    response = client.post(
        "/api/search-runs",
        json={
            "resume_id": resume_id,
            "keywords": ["Python 实习"],
            "city": "上海",
            "platforms": ["boss", "unknown-platform"],
        },
    )

    assert response.status_code == 400
    assert "Unsupported platforms" in response.json()["detail"]


def test_browser_cdp_search_mode_requires_detected_platform_tab(tmp_path, monkeypatch):
    monkeypatch.delenv("BROWSER_CDP_URL", raising=False)
    monkeypatch.setattr(
        platform_session_service,
        "urlopen",
        lambda url, timeout: (_ for _ in ()).throw(OSError("no local CDP")),
    )
    app = create_app(db_path=tmp_path / "agent-business.sqlite3")
    client = TestClient(app)

    upload_response = client.post(
        "/api/resumes",
        files={"file": ("resume.txt", "技能: Python, FastAPI".encode("utf-8"), "text/plain")},
    )
    resume_id = upload_response.json()["id"]

    response = client.post(
        "/api/search-runs",
        json={
            "resume_id": resume_id,
            "keywords": ["Python 实习"],
            "city": "上海",
            "platforms": ["boss"],
            "search_mode": "browser_cdp",
        },
    )

    assert response.status_code == 400
    assert "browser_cdp search requires detected platform tabs" in response.json()["detail"]
    assert client.get("/api/jobs").json() == []
    agent_events = client.get("/api/agent-events").json()
    agents = {agent["agent_name"]: agent for agent in agent_events["agents"]}
    assert agents["JobSearchAgent"]["status"] == "failed"
    assert "browser_cdp search requires detected platform tabs" in agents["JobSearchAgent"]["error"]


def test_failed_orchestrator_task_detail_includes_manual_retry_boundary(tmp_path, monkeypatch):
    monkeypatch.delenv("BROWSER_CDP_URL", raising=False)
    monkeypatch.setattr(
        platform_session_service,
        "urlopen",
        lambda url, timeout: (_ for _ in ()).throw(OSError("no local CDP")),
    )
    app = create_app(db_path=tmp_path / "agent-business.sqlite3")
    client = TestClient(app)

    upload_response = client.post(
        "/api/resumes",
        files={"file": ("resume.txt", "技能: Python, FastAPI".encode("utf-8"), "text/plain")},
    )
    resume_id = upload_response.json()["id"]
    client.post(
        "/api/search-runs",
        json={
            "resume_id": resume_id,
            "keywords": ["Python 实习"],
            "city": "上海",
            "platforms": ["boss"],
            "search_mode": "browser_cdp",
        },
    )
    task_id = client.get("/api/agent-events").json()["orchestrator"]["last_task"]["id"]

    response = client.get(f"/api/orchestrator/tasks/{task_id}")

    assert response.status_code == 200
    retry_suggestion = response.json()["retry_suggestion"]
    assert retry_suggestion["mode"] == "manual_only"
    assert retry_suggestion["retryable"] is True
    assert "CDP" in retry_suggestion["next_action"]
    assert "不会自动投递" in retry_suggestion["safety_boundary"]


def test_browser_cdp_search_mode_saves_extracted_jobs_without_demo_fallback(tmp_path, monkeypatch):
    class FakePlatformSessionService:
        def inspect(self):
            return PlatformSessionsResponse(
                cdp_url="http://127.0.0.1:9222",
                browser_connected=True,
                sessions=[
                    PlatformSession(
                        platform="boss",
                        expected_hosts=["zhipin.com"],
                        state="tab_detected",
                        detected_url="https://www.zhipin.com/web/geek/job",
                        message="已检测到平台标签页",
                    )
                ],
            )

    class FakeBrowserJobExtractorService:
        def extract(self, platforms, limit):
            assert platforms == ["boss"]
            assert limit == 8
            return type(
                "ExtractionResponse",
                (),
                {
                    "extractions": [
                        PlatformJobExtraction(
                            platform="boss",
                            status="success",
                            source_url="https://www.zhipin.com/web/geek/job",
                            jobs=[
                                ExtractedJobCandidate(
                                    platform="boss",
                                    company="真实公司",
                                    title="Python 后端实习生",
                                    city="上海",
                                    salary="200-300/天",
                                    description="真实页面可见岗位内容，需要 Python、FastAPI、SQL。",
                                    url="https://www.zhipin.com/job_detail/abc.html",
                                    job_type="boss_browser",
                                    detail_status="detail_fetched",
                                    detail_reason="详情页已补全岗位要求。",
                                )
                            ],
                        )
                    ]
                },
            )()

    monkeypatch.setattr(job_application_service, "PlatformSessionService", FakePlatformSessionService)
    monkeypatch.setattr(job_application_service, "BrowserJobExtractorService", FakeBrowserJobExtractorService)
    app = create_app(db_path=tmp_path / "agent-business.sqlite3")
    client = TestClient(app)

    upload_response = client.post(
        "/api/resumes",
        files={"file": ("resume.txt", "技能: Python, FastAPI".encode("utf-8"), "text/plain")},
    )
    resume_id = upload_response.json()["id"]

    response = client.post(
        "/api/search-runs",
        json={
            "resume_id": resume_id,
            "keywords": ["Python 实习"],
            "city": "上海",
            "platforms": ["boss"],
            "search_mode": "browser_cdp",
        },
    )

    assert response.status_code == 201
    assert response.json()["status"] == "completed"
    jobs = client.get("/api/jobs").json()
    assert len(jobs) == 1
    assert jobs[0]["platform"] == "boss"
    assert jobs[0]["company"] == "真实公司"
    assert jobs[0]["url"] == "https://www.zhipin.com/job_detail/abc.html"
    assert jobs[0]["detail_status"] == "detail_fetched"
    assert jobs[0]["detail_reason"] == "详情页已补全岗位要求。"
    assert "match" in jobs[0]
    assert "example.test" not in str(jobs)


def test_browser_cdp_search_mode_controls_page_with_keywords_before_extracting(tmp_path, monkeypatch):
    class FakePlatformSessionService:
        def inspect(self):
            return PlatformSessionsResponse(
                cdp_url="http://127.0.0.1:9222",
                browser_connected=True,
                sessions=[
                    PlatformSession(
                        platform="boss",
                        expected_hosts=["zhipin.com"],
                        state="tab_detected",
                        detected_url="https://www.zhipin.com/web/geek/job",
                        message="已检测到平台标签页",
                    )
                ],
            )

    class FakeBrowserJobExtractorService:
        calls = []

        def search_and_extract(self, platforms, keywords, city, limit):
            self.__class__.calls.append(
                {
                    "platforms": platforms,
                    "keywords": keywords,
                    "city": city,
                    "limit": limit,
                }
            )
            return type(
                "ExtractionResponse",
                (),
                {
                    "extractions": [
                        PlatformJobExtraction(
                            platform="boss",
                            status="success",
                            source_url="https://www.zhipin.com/web/geek/job?query=Python",
                            jobs=[
                                ExtractedJobCandidate(
                                    platform="boss",
                                    company="搜索后公司",
                                    title="搜索后 Python 岗位",
                                    city="上海",
                                    salary="18-25K",
                                    description="搜索后真实岗位，需要 Python 和数据分析。",
                                    url="https://www.zhipin.com/job_detail/search.html",
                                    job_type="boss_browser",
                                )
                            ],
                        )
                    ]
                },
            )()

    monkeypatch.setattr(job_application_service, "PlatformSessionService", FakePlatformSessionService)
    monkeypatch.setattr(job_application_service, "BrowserJobExtractorService", FakeBrowserJobExtractorService)
    app = create_app(db_path=tmp_path / "agent-business.sqlite3")
    client = TestClient(app)
    upload_response = client.post(
        "/api/resumes",
        files={"file": ("resume.txt", "技能: Python, 数据分析".encode("utf-8"), "text/plain")},
    )
    resume_id = upload_response.json()["id"]

    response = client.post(
        "/api/search-runs",
        json={
            "resume_id": resume_id,
            "keywords": ["Python 实习", "数据分析"],
            "city": "上海",
            "platforms": ["boss"],
            "search_mode": "browser_cdp",
        },
    )

    assert response.status_code == 201
    assert FakeBrowserJobExtractorService.calls == [
        {
            "platforms": ["boss"],
            "keywords": ["Python 实习", "数据分析"],
            "city": "上海",
            "limit": 8,
        }
    ]
    jobs = client.get(f"/api/jobs?search_run_id={response.json()['id']}").json()
    assert jobs[0]["company"] == "搜索后公司"
    assert jobs[0]["salary"] == "18-25K"


def test_single_job_detail_can_be_refreshed_from_browser_cdp(tmp_path, monkeypatch):
    class FakePlatformSessionService:
        def inspect(self):
            return PlatformSessionsResponse(
                cdp_url="http://127.0.0.1:9222",
                browser_connected=True,
                sessions=[
                    PlatformSession(
                        platform="boss",
                        expected_hosts=["zhipin.com"],
                        state="tab_detected",
                        detected_url="https://www.zhipin.com/web/geek/jobs",
                        message="tab detected",
                    )
                ],
            )

    class FakeBrowserJobExtractorService:
        refreshed = []

        def search_and_extract(self, platforms, keywords, city, limit):
            return type(
                "ExtractionResponse",
                (),
                {
                    "extractions": [
                        PlatformJobExtraction(
                            platform="boss",
                            status="success",
                            source_url="https://www.zhipin.com/web/geek/jobs",
                            jobs=[
                                ExtractedJobCandidate(
                                    platform="boss",
                                    company="Real Company",
                                    title="AI Agent Intern",
                                    city="Shanghai",
                                    salary="薪资读取失败",
                                    description="card only",
                                    url="https://www.zhipin.com/job_detail/refresh.html",
                                    job_type="boss_browser",
                                    detail_status="card_only",
                                    detail_reason="Only card text was available.",
                                )
                            ],
                        )
                    ]
                },
            )()

        def refresh_job_detail(self, platform, url):
            self.__class__.refreshed.append({"platform": platform, "url": url})
            return ExtractedJobCandidate(
                platform=platform,
                company="Real Company",
                title="AI Agent Intern",
                city="Shanghai",
                salary="260-350/天",
                description="完整岗位要求：负责 Agent 工作台、FastAPI、React 和数据分析。",
                url=url,
                job_type="boss_browser",
                detail_status="detail_fetched",
                detail_reason="Detail page refreshed for this job.",
            )

    monkeypatch.setattr(job_application_service, "PlatformSessionService", FakePlatformSessionService)
    monkeypatch.setattr(job_application_service, "BrowserJobExtractorService", FakeBrowserJobExtractorService)
    app = create_app(db_path=tmp_path / "agent-business.sqlite3")
    client = TestClient(app)

    upload_response = client.post(
        "/api/resumes",
        files={"file": ("resume.txt", "Skills: Python, FastAPI, React".encode("utf-8"), "text/plain")},
    )
    resume_id = upload_response.json()["id"]
    run_response = client.post(
        "/api/search-runs",
        json={
            "resume_id": resume_id,
            "keywords": ["Agent Intern"],
            "city": "Shanghai",
            "platforms": ["boss"],
            "search_mode": "browser_cdp",
        },
    )
    job_id = client.get(f"/api/jobs?search_run_id={run_response.json()['id']}").json()[0]["id"]

    response = client.post(f"/api/jobs/{job_id}/refresh-detail")

    assert response.status_code == 200
    refreshed = response.json()
    assert refreshed["id"] == job_id
    assert refreshed["salary"] == "260-350/天"
    assert "完整岗位要求" in refreshed["description"]
    assert refreshed["detail_status"] == "detail_fetched"
    assert refreshed["detail_reason"] == "Detail page refreshed for this job."
    assert FakeBrowserJobExtractorService.refreshed == [
        {"platform": "boss", "url": "https://www.zhipin.com/job_detail/refresh.html"}
    ]
    persisted = client.get(f"/api/jobs?search_run_id={run_response.json()['id']}").json()[0]
    assert persisted["salary"] == "260-350/天"
    assert persisted["detail_status"] == "detail_fetched"

    events = client.get("/api/agent-events").json()
    assert events["orchestrator"]["last_task"]["task_name"] == "job.detail.refresh"
    assert events["orchestrator"]["last_task"]["status"] == "success"
    steps = events["orchestrator"]["last_task"]["steps"]
    assert steps[0]["agent_name"] == "JobSearchAgent"
    assert steps[0]["status"] == "running"
    assert steps[0]["step"] == "refresh job detail"
    assert steps[-1]["status"] == "success"
    assert "detail_fetched" in steps[-1]["output_summary"]


def test_single_job_detail_refresh_failure_records_agent_event(tmp_path):
    app = create_app(db_path=tmp_path / "agent-business.sqlite3")
    client = TestClient(app)

    response = client.post("/api/jobs/404/refresh-detail")

    assert response.status_code == 404
    events = client.get("/api/agent-events").json()
    assert events["orchestrator"]["last_task"]["task_name"] == "job.detail.refresh"
    assert events["orchestrator"]["last_task"]["status"] == "failed"
    failed_step = events["orchestrator"]["last_task"]["steps"][-1]
    assert failed_step["agent_name"] == "JobSearchAgent"
    assert failed_step["status"] == "failed"
    assert "Job 404 not found" in failed_step["error"]


def test_manual_job_detail_update_recalculates_match_and_records_agent_event(tmp_path, monkeypatch):
    class FakePlatformSessionService:
        def inspect(self):
            return PlatformSessionsResponse(
                cdp_url="http://127.0.0.1:9222",
                browser_connected=True,
                sessions=[
                    PlatformSession(
                        platform="boss",
                        expected_hosts=["zhipin.com"],
                        state="tab_detected",
                        detected_url="https://www.zhipin.com/web/geek/jobs",
                        message="tab detected",
                    )
                ],
            )

    class FakeBrowserJobExtractorService:
        def search_and_extract(self, platforms, keywords, city, limit):
            return type(
                "ExtractionResponse",
                (),
                {
                    "extractions": [
                        PlatformJobExtraction(
                            platform="boss",
                            status="success",
                            source_url="https://www.zhipin.com/web/geek/jobs",
                            jobs=[
                                ExtractedJobCandidate(
                                    platform="boss",
                                    company="Manual Detail Co",
                                    title="Operations Intern",
                                    city="Shanghai",
                                    salary="薪资读取失败",
                                    description="整理资料和日常运营支持。",
                                    url="https://www.zhipin.com/job_detail/manual.html",
                                    job_type="boss_browser",
                                    detail_status="card_only",
                                    detail_reason="Only card text was available.",
                                )
                            ],
                        )
                    ]
                },
            )()

    monkeypatch.setattr(job_application_service, "PlatformSessionService", FakePlatformSessionService)
    monkeypatch.setattr(job_application_service, "BrowserJobExtractorService", FakeBrowserJobExtractorService)
    app = create_app(db_path=tmp_path / "agent-business.sqlite3")
    client = TestClient(app)

    upload_response = client.post(
        "/api/resumes",
        files={"file": ("resume.txt", "Skills: Python, FastAPI, React, SQL, Agent".encode("utf-8"), "text/plain")},
    )
    resume_id = upload_response.json()["id"]
    run_response = client.post(
        "/api/search-runs",
        json={
            "resume_id": resume_id,
            "keywords": ["Agent Intern"],
            "city": "Shanghai",
            "platforms": ["boss"],
            "search_mode": "browser_cdp",
        },
    )
    job = client.get(f"/api/jobs?search_run_id={run_response.json()['id']}").json()[0]
    original_score = job["match"]["score"]
    manual_description = "岗位要求：使用 Python、FastAPI、React、SQL 开发 Agent 工作台和数据分析看板。"

    response = client.patch(
        f"/api/jobs/{job['id']}/manual-detail",
        json={"description": manual_description, "note": "用户从平台详情页手动补全 JD"},
    )

    assert response.status_code == 200
    updated = response.json()
    assert updated["description"] == manual_description
    assert updated["detail_status"] == "manual_filled"
    assert "手动补全" in updated["detail_reason"]
    assert updated["match"]["score"] > original_score
    assert "Python" in updated["match"]["hit_reasons"]
    persisted = client.get(f"/api/jobs?search_run_id={run_response.json()['id']}").json()[0]
    assert persisted["match"]["score"] == updated["match"]["score"]

    events = client.get("/api/agent-events").json()
    assert events["orchestrator"]["last_task"]["task_name"] == "job.detail.manual_update"
    assert events["orchestrator"]["last_task"]["status"] == "success"
    steps = events["orchestrator"]["last_task"]["steps"]
    assert steps[0]["agent_name"] == "JobMatchAgent"
    assert steps[0]["status"] == "running"
    assert steps[-1]["status"] == "success"
    assert "score=" in steps[-1]["output_summary"]


def test_tailor_blocks_low_quality_jd_then_uses_manual_detail_for_llm(tmp_path, monkeypatch):
    manual_description = "岗位职责：负责 Agent 工作台自动化评测、提示词实验记录和数据看板。任职要求：熟悉 Python、FastAPI、React。"

    class FakePlatformSessionService:
        def inspect(self):
            return PlatformSessionsResponse(
                cdp_url="http://127.0.0.1:9222",
                browser_connected=True,
                sessions=[
                    PlatformSession(
                        platform="boss",
                        expected_hosts=["zhipin.com"],
                        state="tab_detected",
                        detected_url="https://www.zhipin.com/web/geek/jobs",
                        message="tab detected",
                    )
                ],
            )

    class FakeBrowserJobExtractorService:
        def search_and_extract(self, platforms, keywords, city, limit):
            return type(
                "ExtractionResponse",
                (),
                {
                    "extractions": [
                        PlatformJobExtraction(
                            platform="boss",
                            status="success",
                            source_url="https://www.zhipin.com/web/geek/jobs",
                            jobs=[
                                ExtractedJobCandidate(
                                    platform="boss",
                                    company="Detail Quality Co",
                                    title="Agent 平台开发实习",
                                    city="上海",
                                    salary="薪资读取失败",
                                    description="card only",
                                    url="https://www.zhipin.com/job_detail/detail-quality.html",
                                    job_type="boss_browser",
                                    detail_status="card_only",
                                    detail_reason="当前只读取到列表卡片，详情页未补全。",
                                )
                            ],
                        )
                    ]
                },
            )()

    class FakeLlmResult:
        content = json.dumps(
            {
                "resume_rewrite": "简历改写要求：突出 Agent 工作台自动化评测、提示词实验记录和数据看板经验。",
                "greeting_message": "您好，我对 Agent 平台开发实习很感兴趣，已围绕自动化评测和数据看板准备定制简历。",
                "diff_summary": ["引用补全后的完整 JD"],
                "resume_risk_flags": [],
                "greeting_risk_flags": [],
                "tone": "professional",
            },
            ensure_ascii=False,
        )
        provider = "openai-compatible"
        model = "deepseek-chat"
        prompt_tokens = 90
        completion_tokens = 45
        duration_ms = 20
        estimated = False

    class FakeOpenAICompatibleClient:
        def generate_application_materials(self, config, resume, job):
            assert manual_description in job.description
            assert "card only" not in job.description
            return FakeLlmResult()

    monkeypatch.setattr(job_application_service, "PlatformSessionService", FakePlatformSessionService)
    monkeypatch.setattr(job_application_service, "BrowserJobExtractorService", FakeBrowserJobExtractorService)
    monkeypatch.setattr(job_application_service, "OpenAICompatibleClient", FakeOpenAICompatibleClient, raising=False)
    monkeypatch.setenv("AGENT_BUSINESS_TEST_API_KEY", "secret-value-that-must-not-leak")
    app = create_app(db_path=tmp_path / "agent-business.sqlite3")
    client = TestClient(app)

    upload_response = client.post(
        "/api/resumes",
        files={
            "file": (
                "resume.txt",
                "技能: Python, FastAPI, React\n项目: Agent 工作台，负责 API 和看板。".encode("utf-8"),
                "text/plain",
            )
        },
    )
    resume_id = upload_response.json()["id"]
    run_response = client.post(
        "/api/search-runs",
        json={
            "resume_id": resume_id,
            "keywords": ["Agent 平台开发"],
            "city": "上海",
            "platforms": ["boss"],
            "search_mode": "browser_cdp",
        },
    )
    job = client.get(f"/api/jobs?search_run_id={run_response.json()['id']}").json()[0]

    blocked_response = client.post(f"/api/jobs/{job['id']}/tailor", json={"resume_id": resume_id})

    assert blocked_response.status_code == 409
    assert "补全 JD" in blocked_response.json()["detail"]

    manual_response = client.patch(
        f"/api/jobs/{job['id']}/manual-detail",
        json={"description": manual_description, "note": "人工补全真实岗位详情"},
    )
    assert manual_response.status_code == 200
    client.put(
        "/api/model-config",
        json={
            "provider": "openai-compatible",
            "model": "deepseek-chat",
            "base_url": "https://api.deepseek.com/v1",
            "api_key_env_var": "AGENT_BUSINESS_TEST_API_KEY",
            "enabled": True,
            "estimation_only": False,
            "timeout_ms": 45000,
            "input_price_per_million": 1.0,
            "output_price_per_million": 2.0,
        },
    )

    tailor_response = client.post(f"/api/jobs/{job['id']}/tailor", json={"resume_id": resume_id})

    assert tailor_response.status_code == 201
    payload = tailor_response.json()
    assert "自动化评测" in payload["resume_rewrite"]
    assert "数据看板" in payload["greeting"]["message"]
    assert payload["review"]["llm"]["status"] == "success"


def test_jobs_can_be_filtered_to_the_current_search_run(tmp_path, monkeypatch):
    class FakePlatformSessionService:
        def inspect(self):
            return PlatformSessionsResponse(
                cdp_url="http://127.0.0.1:9222",
                browser_connected=True,
                sessions=[
                    PlatformSession(
                        platform="boss",
                        expected_hosts=["zhipin.com"],
                        state="tab_detected",
                        detected_url="https://www.zhipin.com/web/geek/job",
                        message="已检测到平台标签页",
                    )
                ],
            )

    class FakeBrowserJobExtractorService:
        def extract(self, platforms, limit):
            return type(
                "ExtractionResponse",
                (),
                {
                    "extractions": [
                        PlatformJobExtraction(
                            platform="boss",
                            status="success",
                            source_url="https://www.zhipin.com/web/geek/job",
                            jobs=[
                                ExtractedJobCandidate(
                                    platform="boss",
                                    company="真实公司",
                                    title="真实 Python 岗位",
                                    city="上海",
                                    salary="200-300/天",
                                    description="真实页面岗位，需要 Python。",
                                    url="https://www.zhipin.com/job_detail/real.html",
                                    job_type="boss_browser",
                                )
                            ],
                        )
                    ]
                },
            )()

    monkeypatch.setattr(job_application_service, "PlatformSessionService", FakePlatformSessionService)
    monkeypatch.setattr(job_application_service, "BrowserJobExtractorService", FakeBrowserJobExtractorService)
    app = create_app(db_path=tmp_path / "agent-business.sqlite3")
    client = TestClient(app)

    upload_response = client.post(
        "/api/resumes",
        files={"file": ("resume.txt", "技能: Python, FastAPI".encode("utf-8"), "text/plain")},
    )
    resume_id = upload_response.json()["id"]
    demo_run = client.post(
        "/api/search-runs",
        json={
            "resume_id": resume_id,
            "keywords": ["Demo 实习"],
            "city": "上海",
            "platforms": ["boss"],
            "search_mode": "demo",
        },
    ).json()
    browser_run = client.post(
        "/api/search-runs",
        json={
            "resume_id": resume_id,
            "keywords": ["Python 实习"],
            "city": "上海",
            "platforms": ["boss"],
            "search_mode": "browser_cdp",
        },
    ).json()

    all_jobs = client.get("/api/jobs").json()
    filtered_jobs = client.get(f"/api/jobs?search_run_id={browser_run['id']}").json()

    assert {job["search_run_id"] for job in all_jobs} == {demo_run["id"], browser_run["id"]}
    assert len(filtered_jobs) == 1
    assert filtered_jobs[0]["search_run_id"] == browser_run["id"]
    assert filtered_jobs[0]["company"] == "真实公司"


def test_browser_cdp_search_mode_reports_extraction_failure_without_demo_jobs(tmp_path, monkeypatch):
    class FakePlatformSessionService:
        def inspect(self):
            return PlatformSessionsResponse(
                cdp_url="http://127.0.0.1:9222",
                browser_connected=True,
                sessions=[
                    PlatformSession(
                        platform="boss",
                        expected_hosts=["zhipin.com"],
                        state="tab_detected",
                        detected_url="https://www.zhipin.com/web/geek/job",
                        message="已检测到平台标签页",
                    )
                ],
            )

    class FakeBrowserJobExtractorService:
        def extract(self, platforms, limit):
            return type(
                "ExtractionResponse",
                (),
                {
                    "extractions": [
                        PlatformJobExtraction(
                            platform="boss",
                            status="extract_failed",
                            source_url="https://www.zhipin.com/web/geek/job",
                            error="页面结构变化，未能提取岗位",
                        )
                    ]
                },
            )()

    monkeypatch.setattr(job_application_service, "PlatformSessionService", FakePlatformSessionService)
    monkeypatch.setattr(job_application_service, "BrowserJobExtractorService", FakeBrowserJobExtractorService)
    app = create_app(db_path=tmp_path / "agent-business.sqlite3")
    client = TestClient(app)

    upload_response = client.post(
        "/api/resumes",
        files={"file": ("resume.txt", "技能: Python, FastAPI".encode("utf-8"), "text/plain")},
    )
    resume_id = upload_response.json()["id"]

    response = client.post(
        "/api/search-runs",
        json={
            "resume_id": resume_id,
            "keywords": ["Python 实习"],
            "city": "上海",
            "platforms": ["boss"],
            "search_mode": "browser_cdp",
        },
    )

    assert response.status_code == 400
    assert "browser_cdp extraction failed" in response.json()["detail"]
    assert "页面结构变化" in response.json()["detail"]
    assert client.get("/api/jobs").json() == []
    agent_events = client.get("/api/agent-events").json()
    agents = {agent["agent_name"]: agent for agent in agent_events["agents"]}
    assert agents["JobSearchAgent"]["status"] == "failed"
    assert "browser_cdp extraction failed" in agents["JobSearchAgent"]["error"]


def test_model_config_can_be_saved_without_exposing_api_key(tmp_path, monkeypatch):
    monkeypatch.setenv("AGENT_BUSINESS_TEST_API_KEY", "secret-value-that-must-not-leak")
    app = create_app(db_path=tmp_path / "agent-business.sqlite3")
    client = TestClient(app)

    update_response = client.put(
        "/api/model-config",
        json={
            "provider": "openai-compatible",
            "model": "deepseek-chat",
            "base_url": "https://api.deepseek.com/v1",
            "api_key_env_var": "AGENT_BUSINESS_TEST_API_KEY",
            "enabled": True,
            "estimation_only": False,
            "timeout_ms": 45000,
            "input_price_per_million": 1.0,
            "output_price_per_million": 2.0,
        },
    )

    assert update_response.status_code == 200
    saved = update_response.json()
    assert saved["provider"] == "openai-compatible"
    assert saved["model"] == "deepseek-chat"
    assert saved["api_key_env_var"] == "AGENT_BUSINESS_TEST_API_KEY"
    assert saved["api_key_configured"] is True
    assert "secret-value-that-must-not-leak" not in str(saved)

    get_response = client.get("/api/model-config")

    assert get_response.status_code == 200
    persisted = get_response.json()
    assert persisted["model"] == "deepseek-chat"
    assert persisted["api_key_configured"] is True
    assert "secret-value-that-must-not-leak" not in str(persisted)


def test_model_profiles_can_be_created_updated_applied_and_deleted(tmp_path, monkeypatch):
    monkeypatch.setenv("AGENT_BUSINESS_TEST_API_KEY", "secret-value-that-must-not-leak")
    app = create_app(db_path=tmp_path / "agent-business.sqlite3")
    client = TestClient(app)

    create_response = client.post(
        "/api/model-profiles",
        json={
            "name": "DeepSeek v4pro",
            "provider": "openai-compatible",
            "model": "v4pro",
            "base_url": "https://api.deepseek.com",
            "api_key_env_var": "AGENT_BUSINESS_TEST_API_KEY",
            "enabled": True,
            "estimation_only": False,
            "timeout_ms": 90000,
            "input_price_per_million": 1.0,
            "output_price_per_million": 2.0,
        },
    )

    assert create_response.status_code == 201
    profile = create_response.json()
    assert profile["id"] > 0
    assert profile["model"] == "deepseek-v4-pro"
    assert profile["api_key_configured"] is True
    assert "secret-value-that-must-not-leak" not in str(profile)

    update_response = client.put(
        f"/api/model-profiles/{profile['id']}",
        json={
            "name": "DeepSeek v4flash",
            "provider": "openai-compatible",
            "model": "v4flash",
            "base_url": "https://api.deepseek.com",
            "api_key_env_var": "AGENT_BUSINESS_TEST_API_KEY",
            "enabled": True,
            "estimation_only": False,
            "timeout_ms": 45000,
            "input_price_per_million": 0.5,
            "output_price_per_million": 1.0,
        },
    )

    assert update_response.status_code == 200
    updated = update_response.json()
    assert updated["name"] == "DeepSeek v4flash"
    assert updated["model"] == "deepseek-v4-flash"

    list_response = client.get("/api/model-profiles")

    assert list_response.status_code == 200
    profiles = list_response.json()["profiles"]
    assert [item["id"] for item in profiles] == [profile["id"]]
    assert profiles[0]["api_key_configured"] is True

    apply_response = client.post(f"/api/model-profiles/{profile['id']}/apply")

    assert apply_response.status_code == 200
    applied = apply_response.json()
    assert applied["model"] == "deepseek-v4-flash"
    assert applied["api_key_env_var"] == "AGENT_BUSINESS_TEST_API_KEY"
    assert applied["api_key_configured"] is True

    delete_response = client.delete(f"/api/model-profiles/{profile['id']}")

    assert delete_response.status_code == 204
    assert client.get("/api/model-profiles").json()["profiles"] == []


def test_agent_model_routes_can_be_saved_per_agent(tmp_path, monkeypatch):
    monkeypatch.setenv("AGENT_BUSINESS_TEST_API_KEY", "secret-value-that-must-not-leak")
    app = create_app(db_path=tmp_path / "agent-business.sqlite3")
    client = TestClient(app)

    default_response = client.get("/api/model-routes")

    assert default_response.status_code == 200
    default_routes = {route["agent_name"]: route for route in default_response.json()["routes"]}
    assert {"ApplicationWriterAgent", "JobMatchAgent", "ReviewAgent"}.issubset(default_routes)

    writer_response = client.put(
        "/api/model-routes/ApplicationWriterAgent",
        json={
            "provider": "openai-compatible",
            "model": "v4pro",
            "base_url": "https://api.deepseek.com",
            "api_key_env_var": "AGENT_BUSINESS_TEST_API_KEY",
            "enabled": True,
            "estimation_only": False,
            "timeout_ms": 90000,
            "input_price_per_million": 1.0,
            "output_price_per_million": 2.0,
        },
    )
    matcher_response = client.put(
        "/api/model-routes/JobMatchAgent",
        json={
            "provider": "openai-compatible",
            "model": "v4flash",
            "base_url": "https://api.deepseek.com",
            "api_key_env_var": "AGENT_BUSINESS_TEST_API_KEY",
            "enabled": True,
            "estimation_only": False,
            "timeout_ms": 45000,
            "input_price_per_million": 0.5,
            "output_price_per_million": 1.0,
        },
    )
    review_response = client.put(
        "/api/model-routes/ReviewAgent",
        json={
            "provider": "openai-compatible",
            "model": "deepseek-v4-pro",
            "base_url": "https://api.deepseek.com",
            "api_key_env_var": "AGENT_BUSINESS_TEST_API_KEY",
            "enabled": True,
            "estimation_only": False,
            "timeout_ms": 45000,
            "input_price_per_million": 1.0,
            "output_price_per_million": 2.0,
        },
    )

    assert writer_response.status_code == 200
    assert matcher_response.status_code == 200
    assert review_response.status_code == 200
    saved = {route["agent_name"]: route for route in client.get("/api/model-routes").json()["routes"]}
    assert saved["ApplicationWriterAgent"]["model"] == "deepseek-v4-pro"
    assert saved["JobMatchAgent"]["model"] == "deepseek-v4-flash"
    assert saved["ReviewAgent"]["model"] == "deepseek-v4-pro"
    assert saved["ApplicationWriterAgent"]["api_key_configured"] is True
    assert saved["JobMatchAgent"]["api_key_configured"] is True
    assert saved["ReviewAgent"]["api_key_configured"] is True
    assert "secret-value-that-must-not-leak" not in str(saved)


def test_agent_model_route_rejects_unknown_agent(tmp_path):
    app = create_app(db_path=tmp_path / "agent-business.sqlite3")
    client = TestClient(app)

    response = client.put(
        "/api/model-routes/UnknownAgent",
        json={
            "provider": "openai-compatible",
            "model": "deepseek-v4-pro",
            "base_url": "https://api.deepseek.com",
            "api_key_env_var": "DEEPSEEK_API_KEY",
            "enabled": True,
            "estimation_only": False,
            "timeout_ms": 90000,
            "input_price_per_million": 1.0,
            "output_price_per_million": 2.0,
        },
    )

    assert response.status_code == 400
    assert "Unsupported model route agent" in response.json()["detail"]


def test_default_model_config_uses_deepseek_when_key_env_is_available(tmp_path, monkeypatch):
    monkeypatch.setenv("DEEPSEEK_API_KEY", "secret-value-that-must-not-leak")
    app = create_app(db_path=tmp_path / "agent-business.sqlite3")
    client = TestClient(app)

    response = client.get("/api/model-config")

    assert response.status_code == 200
    payload = response.json()
    assert payload["provider"] == "openai-compatible"
    assert payload["model"] == "deepseek-v4-pro"
    assert payload["base_url"] == "https://api.deepseek.com"
    assert payload["api_key_env_var"] == "DEEPSEEK_API_KEY"
    assert payload["api_key_configured"] is True
    assert payload["enabled"] is True
    assert payload["estimation_only"] is False
    assert "secret-value-that-must-not-leak" not in str(payload)


def test_default_model_config_loads_deepseek_key_from_external_env_file(tmp_path, monkeypatch):
    external_env = tmp_path / "external.env"
    external_env.write_text('DEEPSEEK_API_KEY="secret-value-that-must-not-leak"\n', encoding="utf-8")
    monkeypatch.delenv("DEEPSEEK_API_KEY", raising=False)
    monkeypatch.delenv("OPENAI_API_KEY", raising=False)
    monkeypatch.setenv("AGENT_BUSINESS_ENV_FILE", str(external_env))
    app = create_app(db_path=tmp_path / "agent-business.sqlite3")
    client = TestClient(app)

    response = client.get("/api/model-config")

    assert response.status_code == 200
    payload = response.json()
    assert payload["provider"] == "openai-compatible"
    assert payload["model"] == "deepseek-v4-pro"
    assert payload["base_url"] == "https://api.deepseek.com"
    assert payload["api_key_env_var"] == "DEEPSEEK_API_KEY"
    assert payload["api_key_configured"] is True
    assert payload["enabled"] is True
    assert payload["estimation_only"] is False
    assert "secret-value-that-must-not-leak" not in str(payload)


def test_model_connection_uses_deepseek_key_from_external_env_file(tmp_path, monkeypatch):
    import app.main as main_module

    external_env = tmp_path / "external.env"
    external_env.write_text("DEEPSEEK_API_KEY=secret-value-that-must-not-leak\n", encoding="utf-8")
    monkeypatch.delenv("DEEPSEEK_API_KEY", raising=False)
    monkeypatch.setenv("AGENT_BUSINESS_ENV_FILE", str(external_env))

    class FakeOpenAICompatibleClient:
        def test_connection(self, config):
            assert config.api_key_configured is True
            assert config.model == "deepseek-v4-pro"
            return {
                "status": "success",
                "provider": config.provider,
                "model": config.model,
                "duration_ms": 18,
                "message": "model connection ok",
            }

    monkeypatch.setattr(main_module, "OpenAICompatibleClient", FakeOpenAICompatibleClient)
    app = create_app(db_path=tmp_path / "agent-business.sqlite3")
    client = TestClient(app)

    response = client.post("/api/model-config/test")

    assert response.status_code == 200
    payload = response.json()
    assert payload["status"] == "success"
    assert payload["model"] == "deepseek-v4-pro"
    assert payload["api_key_configured"] is True
    assert "secret-value-that-must-not-leak" not in str(payload)


def test_openai_compatible_client_sends_stable_user_agent(monkeypatch):
    from app.schemas import ModelConfig
    from app.services import llm_client_service
    from app.services.llm_client_service import OpenAICompatibleClient

    captured_headers = {}

    class FakeResponse:
        def __enter__(self):
            return self

        def __exit__(self, exc_type, exc, traceback):
            return False

        def read(self):
            return b'{"choices":[{"message":{"content":"ok"}}]}'

    def fake_urlopen(request, timeout):
        captured_headers.update(dict(request.header_items()))
        return FakeResponse()

    monkeypatch.setenv("AGENT_BUSINESS_TEST_API_KEY", "secret-value-that-must-not-leak")
    monkeypatch.setattr(llm_client_service, "urlopen", fake_urlopen)
    config = ModelConfig(
        provider="openai-compatible",
        model="deepseek-v4-pro",
        base_url="https://api.deepseek.com",
        api_key_env_var="AGENT_BUSINESS_TEST_API_KEY",
        api_key_configured=True,
        enabled=True,
        estimation_only=False,
        timeout_ms=30000,
        input_price_per_million=0,
        output_price_per_million=0,
    )

    OpenAICompatibleClient().test_connection(config)

    assert captured_headers["User-agent"] == "agent-business/0.1"
    assert captured_headers["Accept"] == "application/json"


def test_model_config_connection_test_returns_sanitized_success(tmp_path, monkeypatch):
    import app.main as main_module

    class FakeOpenAICompatibleClient:
        def test_connection(self, config):
            assert config.model == "deepseek-chat"
            assert config.api_key_env_var == "AGENT_BUSINESS_TEST_API_KEY"
            return {
                "status": "success",
                "provider": config.provider,
                "model": config.model,
                "duration_ms": 24,
                "message": "model connection ok",
            }

    monkeypatch.setenv("AGENT_BUSINESS_TEST_API_KEY", "secret-value-that-must-not-leak")
    monkeypatch.setattr(main_module, "OpenAICompatibleClient", FakeOpenAICompatibleClient)
    app = create_app(db_path=tmp_path / "agent-business.sqlite3")
    client = TestClient(app)
    client.put(
        "/api/model-config",
        json={
            "provider": "openai-compatible",
            "model": "deepseek-chat",
            "base_url": "https://api.deepseek.com/v1",
            "api_key_env_var": "AGENT_BUSINESS_TEST_API_KEY",
            "enabled": True,
            "estimation_only": False,
            "timeout_ms": 45000,
            "input_price_per_million": 1.0,
            "output_price_per_million": 2.0,
        },
    )

    response = client.post("/api/model-config/test")

    assert response.status_code == 200
    payload = response.json()
    assert payload["status"] == "success"
    assert payload["provider"] == "openai-compatible"
    assert payload["model"] == "deepseek-chat"
    assert payload["duration_ms"] == 24
    assert payload["api_key_configured"] is True
    assert "secret-value-that-must-not-leak" not in str(payload)


def test_model_config_connection_test_returns_sanitized_failure(tmp_path, monkeypatch):
    import app.main as main_module
    from app.services.llm_client_service import LLMClientUnavailable

    class FakeOpenAICompatibleClient:
        def test_connection(self, config):
            raise LLMClientUnavailable("LLM request failed: HTTPError with secret-value-that-must-not-leak")

    monkeypatch.setenv("AGENT_BUSINESS_TEST_API_KEY", "secret-value-that-must-not-leak")
    monkeypatch.setattr(main_module, "OpenAICompatibleClient", FakeOpenAICompatibleClient)
    app = create_app(db_path=tmp_path / "agent-business.sqlite3")
    client = TestClient(app)
    client.put(
        "/api/model-config",
        json={
            "provider": "openai-compatible",
            "model": "deepseek-chat",
            "base_url": "https://api.deepseek.com/v1",
            "api_key_env_var": "AGENT_BUSINESS_TEST_API_KEY",
            "enabled": True,
            "estimation_only": False,
            "timeout_ms": 45000,
            "input_price_per_million": 1.0,
            "output_price_per_million": 2.0,
        },
    )

    response = client.post("/api/model-config/test")

    assert response.status_code == 200
    payload = response.json()
    assert payload["status"] == "failed"
    assert payload["provider"] == "openai-compatible"
    assert payload["model"] == "deepseek-chat"
    assert payload["api_key_configured"] is True
    assert "HTTPError" in payload["error"]
    assert "secret-value-that-must-not-leak" not in str(payload)


def test_tailor_uses_enabled_openai_compatible_model_and_records_usage(tmp_path, monkeypatch):
    class FakeLlmResult:
        content = json.dumps(
            {
                "resume_text": "LLM 定制简历：保留 Python、FastAPI、React 项目经历。",
                "greeting_message": "LLM 招呼语：您好，我想投递这个后端实习岗位。",
                "diff_summary": ["突出 FastAPI 项目"],
                "resume_risk_flags": [],
                "greeting_risk_flags": [],
                "tone": "professional",
            },
            ensure_ascii=False,
        )
        provider = "openai-compatible"
        model = "deepseek-chat"
        prompt_tokens = 120
        completion_tokens = 60
        duration_ms = 88
        estimated = False

    class FakeOpenAICompatibleClient:
        def generate_application_materials(self, config, resume, job):
            assert config.enabled is True
            assert config.estimation_only is False
            assert config.model == "deepseek-chat"
            return FakeLlmResult()

    monkeypatch.setenv("AGENT_BUSINESS_TEST_API_KEY", "secret-value-that-must-not-leak")
    monkeypatch.setattr(job_application_service, "OpenAICompatibleClient", FakeOpenAICompatibleClient, raising=False)
    db_path = tmp_path / "agent-business.sqlite3"
    app = create_app(db_path=db_path)
    client = TestClient(app)
    client.put(
        "/api/model-config",
        json={
            "provider": "openai-compatible",
            "model": "deepseek-chat",
            "base_url": "https://api.deepseek.com/v1",
            "api_key_env_var": "AGENT_BUSINESS_TEST_API_KEY",
            "enabled": True,
            "estimation_only": False,
            "timeout_ms": 45000,
            "input_price_per_million": 1.0,
            "output_price_per_million": 2.0,
        },
    )

    upload_response = client.post(
        "/api/resumes",
        files={"file": ("resume.txt", "技能: Python, FastAPI, React".encode("utf-8"), "text/plain")},
    )
    resume_id = upload_response.json()["id"]
    client.post(
        "/api/search-runs",
        json={
            "resume_id": resume_id,
            "keywords": ["Python 实习"],
            "city": "上海",
            "platforms": ["boss"],
        },
    )
    job_id = client.get("/api/jobs").json()[0]["id"]

    tailor_response = client.post(f"/api/jobs/{job_id}/tailor", json={"resume_id": resume_id})

    assert tailor_response.status_code == 201
    payload = tailor_response.json()
    assert "LLM 定制简历" in payload["resume_text"]
    assert "LLM 招呼语" in payload["greeting"]["message"]
    assert payload["review"]["llm"]["status"] == "success"
    assert payload["review"]["llm"]["prompt_tokens"] == 120
    assert payload["review"]["llm"]["completion_tokens"] == 60
    assert payload["review"]["llm"]["total_tokens"] == 180
    assert payload["review"]["llm"]["cost_usd"] == 0.00024
    assert payload["review"]["llm"]["duration_ms"] == 88
    assert payload["review"]["llm"]["usage_status"] == "success"
    assert payload["review"]["llm"]["usage_estimated"] is False
    assert "secret-value-that-must-not-leak" not in str(payload)
    with sqlite3.connect(db_path) as conn:
        row = conn.execute(
            "SELECT provider, model, prompt_tokens, completion_tokens, estimated, cost_usd "
            "FROM llm_usage WHERE agent_name = 'ApplicationWriterAgent' ORDER BY id DESC LIMIT 1"
        ).fetchone()
    assert row == ("openai-compatible", "deepseek-chat", 120, 60, 0, 0.00024)


def test_tailor_uses_application_writer_agent_model_route_over_global_config(tmp_path, monkeypatch):
    class FakeLlmResult:
        content = json.dumps(
            {
                "resume_text": "Agent route resume from flash model.",
                "greeting_message": "Agent route greeting from flash model.",
                "diff_summary": ["used ApplicationWriterAgent route"],
                "resume_risk_flags": [],
                "greeting_risk_flags": [],
                "tone": "professional",
            },
            ensure_ascii=False,
        )
        provider = "openai-compatible"
        model = "deepseek-v4-flash"
        prompt_tokens = 80
        completion_tokens = 40
        duration_ms = 44
        estimated = False

    class FakeOpenAICompatibleClient:
        def generate_application_materials(self, config, resume, job):
            assert config.model == "deepseek-v4-flash"
            assert config.timeout_ms == 45000
            return FakeLlmResult()

    monkeypatch.setenv("AGENT_BUSINESS_TEST_API_KEY", "secret-value-that-must-not-leak")
    monkeypatch.setattr(job_application_service, "OpenAICompatibleClient", FakeOpenAICompatibleClient, raising=False)
    db_path = tmp_path / "agent-business.sqlite3"
    app = create_app(db_path=db_path)
    client = TestClient(app)
    client.put(
        "/api/model-config",
        json={
            "provider": "openai-compatible",
            "model": "deepseek-chat",
            "base_url": "https://api.deepseek.com",
            "api_key_env_var": "AGENT_BUSINESS_TEST_API_KEY",
            "enabled": True,
            "estimation_only": False,
            "timeout_ms": 90000,
            "input_price_per_million": 1.0,
            "output_price_per_million": 2.0,
        },
    )
    client.put(
        "/api/model-routes/ApplicationWriterAgent",
        json={
            "provider": "openai-compatible",
            "model": "deepseek-v4-flash",
            "base_url": "https://api.deepseek.com",
            "api_key_env_var": "AGENT_BUSINESS_TEST_API_KEY",
            "enabled": True,
            "estimation_only": False,
            "timeout_ms": 45000,
            "input_price_per_million": 0.5,
            "output_price_per_million": 1.0,
        },
    )

    upload_response = client.post(
        "/api/resumes",
        files={"file": ("resume.txt", "技能: Python, FastAPI, React".encode("utf-8"), "text/plain")},
    )
    resume_id = upload_response.json()["id"]
    client.post(
        "/api/search-runs",
        json={
            "resume_id": resume_id,
            "keywords": ["Python 实习"],
            "city": "上海",
            "platforms": ["boss"],
        },
    )
    job_id = client.get("/api/jobs").json()[0]["id"]

    response = client.post(f"/api/jobs/{job_id}/tailor", json={"resume_id": resume_id})

    assert response.status_code == 201
    payload = response.json()
    assert payload["review"]["llm"]["model"] == "deepseek-v4-flash"
    assert payload["review"]["llm"]["route"]["model"] == "deepseek-v4-flash"
    assert payload["review"]["llm"]["total_tokens"] == 120
    assert "deepseek-chat" not in str(payload["review"]["llm"])
    with sqlite3.connect(db_path) as conn:
        row = conn.execute(
            "SELECT model, prompt_tokens, completion_tokens, cost_usd "
            "FROM llm_usage WHERE agent_name = 'ApplicationWriterAgent' ORDER BY id DESC LIMIT 1"
        ).fetchone()
    assert row == ("deepseek-v4-flash", 80, 40, 0.00008)


def test_search_run_uses_job_match_agent_model_route_for_batch_scoring(tmp_path, monkeypatch):
    class FakeLlmResult:
        content = json.dumps(
            {
                "matches": [
                    {
                        "job_index": 0,
                        "score": 91,
                        "hit_reasons": ["LLM scored Python and Agent fit"],
                        "gap_reasons": ["LLM saw SQL depth gap"],
                        "recommendation": "strong_apply",
                    },
                ]
            },
            ensure_ascii=False,
        )
        provider = "openai-compatible"
        model = "deepseek-v4-flash"
        prompt_tokens = 100
        completion_tokens = 50
        duration_ms = 33
        estimated = False

    class FakeOpenAICompatibleClient:
        def score_job_matches(self, config, resume, jobs, rule_matches):
            assert config.model == "deepseek-v4-flash"
            assert config.timeout_ms == 30000
            assert len(jobs) >= 1
            assert all(match.score >= 0 for match in rule_matches)
            return FakeLlmResult()

    monkeypatch.setenv("AGENT_BUSINESS_TEST_API_KEY", "secret-value-that-must-not-leak")
    monkeypatch.setattr(job_application_service, "OpenAICompatibleClient", FakeOpenAICompatibleClient, raising=False)
    db_path = tmp_path / "agent-business.sqlite3"
    app = create_app(db_path=db_path)
    client = TestClient(app)
    client.put(
        "/api/model-routes/JobMatchAgent",
        json={
            "provider": "openai-compatible",
            "model": "deepseek-v4-flash",
            "base_url": "https://api.deepseek.com",
            "api_key_env_var": "AGENT_BUSINESS_TEST_API_KEY",
            "enabled": True,
            "estimation_only": False,
            "timeout_ms": 30000,
            "input_price_per_million": 0.2,
            "output_price_per_million": 0.4,
        },
    )
    upload_response = client.post(
        "/api/resumes",
        files={
            "file": (
                "resume.txt",
                "技能: Python, FastAPI, React, Agent, 数据分析".encode("utf-8"),
                "text/plain",
            )
        },
    )
    resume_id = upload_response.json()["id"]

    response = client.post(
        "/api/search-runs",
        json={
            "resume_id": resume_id,
            "keywords": ["Python 实习", "Agent 实习"],
            "city": "上海",
            "platforms": ["boss"],
        },
    )

    assert response.status_code == 201
    jobs = sorted(client.get("/api/jobs").json(), key=lambda item: item["match"]["score"], reverse=True)
    assert jobs[0]["match"]["score"] == 91
    assert jobs[0]["match"]["recommendation"] == "strong_apply"
    assert "LLM scored Python and Agent fit" in jobs[0]["match"]["hit_reasons"]
    with sqlite3.connect(db_path) as conn:
        row = conn.execute(
            "SELECT model, prompt_tokens, completion_tokens, cost_usd "
            "FROM llm_usage WHERE agent_name = 'JobMatchAgent' ORDER BY id DESC LIMIT 1"
        ).fetchone()
    assert row == ("deepseek-v4-flash", 100, 50, 0.00004)


def test_tailored_resume_pdf_can_be_downloaded_after_material_generation(tmp_path, monkeypatch):
    class FakeTailoredResumePdfService:
        def render(self, tailored_bundle, template_bytes):
            assert template_bytes.startswith(b"PK")
            assert tailored_bundle["project_rewrite"]
            return _one_page_pdf()

    import app.main as main_module

    monkeypatch.setattr(main_module, "TailoredResumePdfService", FakeTailoredResumePdfService)
    from docx import Document

    document = Document()
    document.add_paragraph("胡俊")
    document.add_paragraph("项目经历")
    document.add_paragraph("旧项目描述")
    buffer = BytesIO()
    document.save(buffer)
    app = create_app(db_path=tmp_path / "agent-business.sqlite3")
    client = TestClient(app)
    upload_response = client.post(
        "/api/resumes",
        files={
            "file": (
                "resume.docx",
                buffer.getvalue(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    resume_id = upload_response.json()["id"]
    client.post(
        "/api/search-runs",
        json={
            "resume_id": resume_id,
            "keywords": ["Python 实习"],
            "city": "上海",
            "platforms": ["boss"],
        },
    )
    job_id = client.get("/api/jobs").json()[0]["id"]
    tailor_response = client.post(f"/api/jobs/{job_id}/tailor", json={"resume_id": resume_id})
    tailored_id = tailor_response.json()["id"]

    pdf_response = client.get(f"/api/tailored-resumes/{tailored_id}/pdf")

    assert pdf_response.status_code == 200
    assert pdf_response.headers["content-type"] == "application/pdf"
    assert pdf_response.content.startswith(b"%PDF")
    assert str(tailored_id) in pdf_response.headers["content-disposition"]
    assert client.get("/api/tailored-resumes/999999/pdf").status_code == 404


def test_tailored_resume_pdf_requires_docx_template(tmp_path):
    app = create_app(db_path=tmp_path / "agent-business.sqlite3")
    client = TestClient(app)
    upload_response = client.post(
        "/api/resumes",
        files={"file": ("resume.txt", "技能: Python, FastAPI, React".encode("utf-8"), "text/plain")},
    )
    resume_id = upload_response.json()["id"]
    client.post(
        "/api/search-runs",
        json={
            "resume_id": resume_id,
            "keywords": ["Python 实习"],
            "city": "上海",
            "platforms": ["boss"],
        },
    )
    job_id = client.get("/api/jobs").json()[0]["id"]
    tailor_response = client.post(f"/api/jobs/{job_id}/tailor", json={"resume_id": resume_id})
    tailored_id = tailor_response.json()["id"]

    pdf_response = client.get(f"/api/tailored-resumes/{tailored_id}/pdf")

    assert pdf_response.status_code == 409
    assert "重新上传 DOCX" in pdf_response.json()["detail"]


def test_tailor_routes_application_writer_through_model_router(tmp_path, monkeypatch):
    from app.services.model_router_service import ModelRoute

    class FakeLlmResult:
        content = json.dumps(
            {
                "resume_text": "Router LLM resume keeps Python and FastAPI.",
                "greeting_message": "Router LLM greeting.",
                "diff_summary": [],
                "resume_risk_flags": [],
                "greeting_risk_flags": [],
                "tone": "professional",
            },
        )
        provider = "openai-compatible"
        model = "deepseek-chat"
        prompt_tokens = 30
        completion_tokens = 12
        duration_ms = 40
        estimated = False

    class FakeModelRouterService:
        requested_agents = []
        requested_configs = []

        def route_for_agent(self, agent_name, config):
            self.requested_agents.append(agent_name)
            self.requested_configs.append((agent_name, config.model))
            if agent_name == "ApplicationWriterAgent":
                return ModelRoute(
                    agent_name=agent_name,
                    mode="external",
                    provider=config.provider,
                    model=config.model,
                    reason="application writing uses configured external model",
                )
            if agent_name == "ReviewAgent":
                return ModelRoute(
                    agent_name=agent_name,
                    mode="external",
                    provider=config.provider,
                    model=config.model,
                    reason="review uses configured external model route",
                )
            return ModelRoute(
                agent_name=agent_name,
                mode="local_rule",
                provider="local",
                model="local-rule",
                reason="review stays local in 4C",
            )

    class FakeOpenAICompatibleClient:
        def generate_application_materials(self, config, resume, job):
            return FakeLlmResult()

    monkeypatch.setenv("AGENT_BUSINESS_TEST_API_KEY", "secret-value-that-must-not-leak")
    monkeypatch.setattr(job_application_service, "ModelRouterService", FakeModelRouterService, raising=False)
    monkeypatch.setattr(job_application_service, "OpenAICompatibleClient", FakeOpenAICompatibleClient, raising=False)
    app = create_app(db_path=tmp_path / "agent-business.sqlite3")
    client = TestClient(app)
    client.put(
        "/api/model-config",
        json={
            "provider": "openai-compatible",
            "model": "deepseek-chat",
            "base_url": "https://api.deepseek.com/v1",
            "api_key_env_var": "AGENT_BUSINESS_TEST_API_KEY",
            "enabled": True,
            "estimation_only": False,
            "timeout_ms": 45000,
            "input_price_per_million": 1.0,
            "output_price_per_million": 2.0,
        },
    )
    client.put(
        "/api/model-routes/ReviewAgent",
        json={
            "provider": "openai-compatible",
            "model": "deepseek-v4-review",
            "base_url": "https://api.deepseek.com/v1",
            "api_key_env_var": "AGENT_BUSINESS_TEST_API_KEY",
            "enabled": True,
            "estimation_only": False,
            "timeout_ms": 45000,
            "input_price_per_million": 1.0,
            "output_price_per_million": 2.0,
        },
    )
    upload_response = client.post(
        "/api/resumes",
        files={"file": ("resume.txt", "Skills: Python, FastAPI, React".encode("utf-8"), "text/plain")},
    )
    resume_id = upload_response.json()["id"]
    client.post(
        "/api/search-runs",
        json={
            "resume_id": resume_id,
            "keywords": ["Python intern"],
            "city": "Shanghai",
            "platforms": ["boss"],
        },
    )
    job_id = client.get("/api/jobs").json()[0]["id"]

    response = client.post(f"/api/jobs/{job_id}/tailor", json={"resume_id": resume_id})

    assert response.status_code == 201
    payload = response.json()
    assert payload["review"]["llm"]["route"]["mode"] == "external"
    assert payload["review"]["llm"]["review_route"]["mode"] == "external"
    assert payload["review"]["llm"]["review_route"]["model"] == "deepseek-v4-review"
    assert FakeModelRouterService.requested_agents == ["JobMatchAgent", "ApplicationWriterAgent", "ReviewAgent"]
    assert ("ReviewAgent", "deepseek-v4-review") in FakeModelRouterService.requested_configs


def test_agent_events_endpoint_reports_real_backend_steps(tmp_path, monkeypatch):
    monkeypatch.delenv("DEEPSEEK_API_KEY", raising=False)
    monkeypatch.delenv("OPENAI_API_KEY", raising=False)
    monkeypatch.delenv("OPENAI_API_KEY_ENV_VAR", raising=False)
    monkeypatch.delenv("LLM_API_KEY_ENV_VAR", raising=False)
    monkeypatch.delenv("LLM_ENABLED", raising=False)
    monkeypatch.delenv("LLM_ESTIMATION_ONLY", raising=False)
    monkeypatch.setenv("AGENT_BUSINESS_ENV_FILE", str(tmp_path / "missing.env"))
    app = create_app(db_path=tmp_path / "agent-business.sqlite3")
    client = TestClient(app)

    upload_response = client.post(
        "/api/resumes",
        files={"file": ("resume.txt", "Skills: Python, FastAPI, React".encode("utf-8"), "text/plain")},
    )
    resume_id = upload_response.json()["id"]
    client.post(
        "/api/search-runs",
        json={
            "resume_id": resume_id,
            "keywords": ["Python intern"],
            "city": "Shanghai",
            "platforms": ["boss"],
        },
    )
    job_id = client.get("/api/jobs").json()[0]["id"]
    client.post(f"/api/jobs/{job_id}/tailor", json={"resume_id": resume_id})

    response = client.get("/api/agent-events")

    assert response.status_code == 200
    payload = response.json()
    agents = {agent["agent_name"]: agent for agent in payload["agents"]}
    assert payload["current_running_agent"] is None
    assert agents["ResumeParserAgent"]["status"] == "success"
    assert agents["JobSearchAgent"]["status"] == "success"
    assert agents["JobMatchAgent"]["status"] == "success"
    assert agents["ApplicationWriterAgent"]["status"] == "success"
    assert agents["ReviewAgent"]["status"] == "success"
    assert agents["ApplicationWriterAgent"]["input_summary"]
    assert agents["ReviewAgent"]["output_summary"]
    assert payload["total_cost_usd"] >= 0
    orchestrator = payload["orchestrator"]
    assert orchestrator["current_task_id"] is None
    assert orchestrator["last_task"]["task_name"] == "application.materials"
    assert orchestrator["last_task"]["status"] == "success"
    assert [
        f"{step['agent_name']}:{step['status']}"
        for step in orchestrator["last_task"]["steps"]
    ] == [
        "ApplicationWriterAgent:running",
        "ApplicationWriterAgent:success",
        "ReviewAgent:running",
        "ReviewAgent:success",
    ]


def test_orchestrator_task_summary_survives_service_restart(tmp_path):
    db_path = tmp_path / "agent-business.sqlite3"
    app = create_app(db_path=db_path)
    client = TestClient(app)

    upload_response = client.post(
        "/api/resumes",
        files={"file": ("resume.txt", "Skills: Python, FastAPI, React".encode("utf-8"), "text/plain")},
    )
    resume_id = upload_response.json()["id"]
    client.post(
        "/api/search-runs",
        json={
            "resume_id": resume_id,
            "keywords": ["Python intern"],
            "city": "Shanghai",
            "platforms": ["boss"],
        },
    )
    job_id = client.get("/api/jobs").json()[0]["id"]
    client.post(f"/api/jobs/{job_id}/tailor", json={"resume_id": resume_id})

    restarted_app = create_app(db_path=db_path)
    restarted_client = TestClient(restarted_app)
    response = restarted_client.get("/api/agent-events")

    assert response.status_code == 200
    orchestrator = response.json()["orchestrator"]
    assert orchestrator["current_task_id"] is None
    assert orchestrator["last_task"]["task_name"] == "application.materials"
    assert orchestrator["last_task"]["status"] == "success"
    assert [step["agent_name"] for step in orchestrator["last_task"]["steps"]] == [
        "ApplicationWriterAgent",
        "ApplicationWriterAgent",
        "ReviewAgent",
        "ReviewAgent",
    ]


def test_orchestrator_task_detail_endpoint_returns_steps(tmp_path):
    db_path = tmp_path / "agent-business.sqlite3"
    app = create_app(db_path=db_path)
    client = TestClient(app)

    upload_response = client.post(
        "/api/resumes",
        files={"file": ("resume.txt", "Skills: Python, FastAPI, React".encode("utf-8"), "text/plain")},
    )
    resume_id = upload_response.json()["id"]
    client.post(
        "/api/search-runs",
        json={
            "resume_id": resume_id,
            "keywords": ["Python intern"],
            "city": "Shanghai",
            "platforms": ["boss"],
        },
    )
    job_id = client.get("/api/jobs").json()[0]["id"]
    client.post(f"/api/jobs/{job_id}/tailor", json={"resume_id": resume_id})
    task_id = client.get("/api/agent-events").json()["orchestrator"]["last_task"]["id"]

    response = client.get(f"/api/orchestrator/tasks/{task_id}")

    assert response.status_code == 200
    payload = response.json()
    assert payload["id"] == task_id
    assert payload["task_name"] == "application.materials"
    assert payload["status"] == "success"
    assert [step["agent_name"] for step in payload["steps"]] == [
        "ApplicationWriterAgent",
        "ApplicationWriterAgent",
        "ReviewAgent",
        "ReviewAgent",
    ]

    missing_response = client.get("/api/orchestrator/tasks/999999")
    assert missing_response.status_code == 404


def test_tailor_falls_back_locally_when_openai_compatible_model_fails(tmp_path, monkeypatch):
    class FakeOpenAICompatibleClient:
        def generate_application_materials(self, config, resume, job):
            raise RuntimeError("upstream timeout")

    monkeypatch.setenv("AGENT_BUSINESS_TEST_API_KEY", "secret-value-that-must-not-leak")
    monkeypatch.setattr(job_application_service, "OpenAICompatibleClient", FakeOpenAICompatibleClient, raising=False)
    db_path = tmp_path / "agent-business.sqlite3"
    app = create_app(db_path=db_path)
    client = TestClient(app)
    client.put(
        "/api/model-config",
        json={
            "provider": "openai-compatible",
            "model": "deepseek-chat",
            "base_url": "https://api.deepseek.com/v1",
            "api_key_env_var": "AGENT_BUSINESS_TEST_API_KEY",
            "enabled": True,
            "estimation_only": False,
            "timeout_ms": 45000,
            "input_price_per_million": 1.0,
            "output_price_per_million": 2.0,
        },
    )

    upload_response = client.post(
        "/api/resumes",
        files={"file": ("resume.txt", "技能: Python, FastAPI, React".encode("utf-8"), "text/plain")},
    )
    resume_id = upload_response.json()["id"]
    client.post(
        "/api/search-runs",
        json={
            "resume_id": resume_id,
            "keywords": ["Python 实习"],
            "city": "上海",
            "platforms": ["boss"],
        },
    )
    job_id = client.get("/api/jobs").json()[0]["id"]

    tailor_response = client.post(f"/api/jobs/{job_id}/tailor", json={"resume_id": resume_id})

    assert tailor_response.status_code == 201
    payload = tailor_response.json()
    assert payload["greeting"]["message"]
    assert payload["review"]["llm"]["status"] == "fallback"
    assert "upstream timeout" in payload["review"]["llm"]["error"]
    assert "secret-value-that-must-not-leak" not in str(payload)
    agent_events = client.get("/api/agent-events").json()
    agents = {agent["agent_name"]: agent for agent in agent_events["agents"]}
    assert agents["ApplicationWriterAgent"]["status"] == "failed"
    assert "upstream timeout" in agents["ApplicationWriterAgent"]["error"]
    usage = client.get("/api/metrics/llm-usage").json()
    writer_bucket = usage["by_agent"]["ApplicationWriterAgent"]
    assert writer_bucket["failed_calls"] == 1
    with sqlite3.connect(db_path) as conn:
        row = conn.execute(
            "SELECT status, error, total_tokens FROM llm_usage "
            "WHERE agent_name = 'ApplicationWriterAgent' AND status = 'failed' "
            "ORDER BY id DESC LIMIT 1"
        ).fetchone()
    assert row[0] == "failed"
    assert "upstream timeout" in row[1]
    assert row[2] == 0


def test_llm_usage_summary_exposes_status_buckets_for_estimated_and_success_calls(tmp_path, monkeypatch):
    class FakeLlmResult:
        content = json.dumps(
            {
                "resume_text": "LLM 定制简历：保留 Python、FastAPI 项目。",
                "greeting_message": "LLM 招呼语：您好。",
                "diff_summary": [],
                "resume_risk_flags": [],
                "greeting_risk_flags": [],
                "tone": "professional",
            },
            ensure_ascii=False,
        )
        provider = "openai-compatible"
        model = "deepseek-chat"
        prompt_tokens = 10
        completion_tokens = 5
        duration_ms = 25
        estimated = False

    class FakeOpenAICompatibleClient:
        def generate_application_materials(self, config, resume, job):
            return FakeLlmResult()

    monkeypatch.setenv("AGENT_BUSINESS_TEST_API_KEY", "secret-value-that-must-not-leak")
    monkeypatch.setattr(job_application_service, "OpenAICompatibleClient", FakeOpenAICompatibleClient, raising=False)
    app = create_app(db_path=tmp_path / "agent-business.sqlite3")
    client = TestClient(app)
    client.put(
        "/api/model-config",
        json={
            "provider": "openai-compatible",
            "model": "deepseek-chat",
            "base_url": "https://api.deepseek.com/v1",
            "api_key_env_var": "AGENT_BUSINESS_TEST_API_KEY",
            "enabled": True,
            "estimation_only": False,
            "timeout_ms": 45000,
            "input_price_per_million": 1.0,
            "output_price_per_million": 2.0,
        },
    )
    upload_response = client.post(
        "/api/resumes",
        files={"file": ("resume.txt", "技能: Python, FastAPI, React".encode("utf-8"), "text/plain")},
    )
    resume_id = upload_response.json()["id"]
    client.post(
        "/api/search-runs",
        json={
            "resume_id": resume_id,
            "keywords": ["Python 实习"],
            "city": "上海",
            "platforms": ["boss"],
        },
    )
    job_id = client.get("/api/jobs").json()[0]["id"]
    client.post(f"/api/jobs/{job_id}/tailor", json={"resume_id": resume_id})

    usage = client.get("/api/metrics/llm-usage").json()

    assert usage["by_agent"]["ApplicationWriterAgent"]["success_calls"] == 1
    assert usage["by_agent"]["ResumeParserAgent"]["estimated_calls"] == 1


def test_platform_sessions_report_missing_cdp_configuration(tmp_path, monkeypatch):
    monkeypatch.delenv("BROWSER_CDP_URL", raising=False)
    monkeypatch.setattr(
        platform_session_service,
        "urlopen",
        lambda url, timeout: (_ for _ in ()).throw(OSError("no local CDP")),
    )
    app = create_app(db_path=tmp_path / "agent-business.sqlite3")
    client = TestClient(app)

    response = client.get("/api/platform-sessions")

    assert response.status_code == 200
    payload = response.json()
    assert payload["browser_connected"] is False
    assert payload["cdp_url"] is None
    assert {session["state"] for session in payload["sessions"]} == {"not_configured"}


def test_platform_sessions_reuses_default_local_cdp_when_env_is_missing(tmp_path, monkeypatch):
    class FakeCdpResponse:
        def __enter__(self):
            return self

        def __exit__(self, exc_type, exc, traceback):
            return False

        def read(self):
            return json.dumps(
                [
                    {
                        "type": "page",
                        "url": "https://www.zhipin.com/web/geek/jobs?query=Python",
                    }
                ]
            ).encode("utf-8")

    def fake_urlopen(url, timeout):
        assert url == "http://127.0.0.1:9222/json"
        return FakeCdpResponse()

    monkeypatch.delenv("BROWSER_CDP_URL", raising=False)
    monkeypatch.setattr(platform_session_service, "urlopen", fake_urlopen)
    app = create_app(db_path=tmp_path / "agent-business.sqlite3")
    client = TestClient(app)

    response = client.get("/api/platform-sessions")

    assert response.status_code == 200
    payload = response.json()
    sessions = {session["platform"]: session for session in payload["sessions"]}
    assert payload["browser_connected"] is True
    assert payload["cdp_url"] == "http://127.0.0.1:9222"
    assert sessions["boss"]["state"] == "tab_detected"
    assert sessions["shixiseng"]["state"] == "tab_not_found"


def test_cdp_browser_launcher_reuses_existing_session_without_starting_process(monkeypatch):
    class FakeCdpResponse:
        def __enter__(self):
            return self

        def __exit__(self, exc_type, exc, tb):
            return False

        def read(self):
            return json.dumps([{"url": "https://www.zhipin.com/web/geek/jobs"}]).encode("utf-8")

    def fail_find_browser(self):
        raise AssertionError("existing CDP session should be reused before finding a browser executable")

    monkeypatch.setenv("BROWSER_CDP_URL", "127.0.0.1:9222")
    monkeypatch.setattr(platform_session_service, "urlopen", lambda url, timeout: FakeCdpResponse())
    monkeypatch.setattr(platform_session_service.CdpBrowserLauncher, "_find_browser", fail_find_browser)

    result = platform_session_service.CdpBrowserLauncher().launch()

    assert result["status"] == "reused"
    assert result["cdp_url"] == "127.0.0.1:9222"
    assert result["opened_urls"] == []


def test_platform_sessions_detect_open_recruitment_tabs(tmp_path, monkeypatch):
    class FakeCdpResponse:
        def __enter__(self):
            return self

        def __exit__(self, exc_type, exc, traceback):
            return False

        def read(self):
            return json.dumps(
                [
                    {
                        "type": "page",
                        "url": "https://www.zhipin.com/web/geek/job?query=Python&token=hidden",
                    },
                    {
                        "type": "page",
                        "url": "https://example.com/",
                    },
                ]
            ).encode("utf-8")

    def fake_urlopen(url, timeout):
        assert url == "http://127.0.0.1:9222/json"
        assert timeout == 2
        return FakeCdpResponse()

    monkeypatch.setenv("BROWSER_CDP_URL", "127.0.0.1:9222")
    monkeypatch.setattr(platform_session_service, "urlopen", fake_urlopen)
    app = create_app(db_path=tmp_path / "agent-business.sqlite3")
    client = TestClient(app)

    response = client.get("/api/platform-sessions")

    assert response.status_code == 200
    payload = response.json()
    sessions = {session["platform"]: session for session in payload["sessions"]}
    assert payload["browser_connected"] is True
    assert sessions["boss"]["state"] == "tab_detected"
    assert sessions["boss"]["detected_url"] == "https://www.zhipin.com/web/geek/job"
    assert sessions["boss"]["authenticated"] is None
    assert sessions["shixiseng"]["state"] == "tab_not_found"
    assert "token=hidden" not in str(payload)


def test_platform_job_extraction_reads_visible_jobs_from_detected_browser_tab(tmp_path, monkeypatch):
    class FakeCdpResponse:
        def __enter__(self):
            return self

        def __exit__(self, exc_type, exc, traceback):
            return False

        def read(self):
            return json.dumps(
                [
                    {
                        "type": "page",
                        "url": "https://www.zhipin.com/web/geek/job?query=Python&token=hidden",
                        "webSocketDebuggerUrl": "ws://127.0.0.1:9222/devtools/page/1",
                    }
                ]
            ).encode("utf-8")

    def fake_urlopen(url, timeout):
        assert url == "http://127.0.0.1:9222/json"
        assert timeout == 2
        return FakeCdpResponse()

    def fake_evaluate(self, websocket_url, expression):
        assert websocket_url == "ws://127.0.0.1:9222/devtools/page/1"
        assert "document.querySelectorAll" in expression
        return {
            "jobs": [
                {
                    "title": "Python 后端实习生",
                    "company": "真实公司",
                    "city": "上海",
                    "salary": "200-300/天",
                    "description": "真实页面可见岗位内容",
                    "url": "https://www.zhipin.com/job_detail/abc.html?securityId=hidden",
                    "job_type": "backend",
                }
            ],
            "diagnostics": {
                "matched_selector_counts": {".job-card-wrapper": 1},
                "candidate_card_count": 1,
            },
        }

    monkeypatch.setenv("BROWSER_CDP_URL", "127.0.0.1:9222")
    monkeypatch.setattr(browser_job_extractor_service, "urlopen", fake_urlopen)
    monkeypatch.setattr(browser_job_extractor_service.CdpRuntimeClient, "evaluate", fake_evaluate)
    app = create_app(db_path=tmp_path / "agent-business.sqlite3")
    client = TestClient(app)

    response = client.post("/api/platform-jobs/extract", json={"platforms": ["boss"], "limit": 5})

    assert response.status_code == 200
    payload = response.json()
    extraction = payload["extractions"][0]
    assert extraction["platform"] == "boss"
    assert extraction["status"] == "success"
    assert extraction["source_url"] == "https://www.zhipin.com/web/geek/job"
    assert extraction["jobs"][0]["company"] == "真实公司"
    assert extraction["jobs"][0]["url"] == "https://www.zhipin.com/job_detail/abc.html"
    assert extraction["diagnostics"]["tab_detected"] is True
    assert extraction["diagnostics"]["websocket_detected"] is True
    assert extraction["diagnostics"]["candidate_card_count"] == 1
    assert extraction["diagnostics"]["extracted_job_count"] == 1
    assert extraction["diagnostics"]["matched_selector_counts"][".job-card-wrapper"] == 1
    assert "hidden" not in str(payload)


def test_platform_job_search_controls_page_before_extracting(tmp_path, monkeypatch):
    class FakeCdpResponse:
        def __enter__(self):
            return self

        def __exit__(self, exc_type, exc, traceback):
            return False

        def read(self):
            return json.dumps(
                [
                    {
                        "type": "page",
                        "url": "https://www.zhipin.com/web/geek/job",
                        "webSocketDebuggerUrl": "ws://127.0.0.1:9222/devtools/page/1",
                    }
                ]
            ).encode("utf-8")

    expressions = []

    def fake_urlopen(url, timeout=2):
        return FakeCdpResponse()

    def fake_evaluate(self, websocket_url, expression):
        expressions.append(expression)
        if "searchUrl" in expression:
            return {"clicked": True, "keyword": "Python 实习", "city": "上海"}
        return {
            "jobs": [
                {
                    "title": "杭州 Python 岗位",
                    "company": "错误城市公司",
                    "city": "杭州",
                    "salary": "18-25K",
                    "description": "错误城市岗位，不应进入上海搜索结果。",
                    "url": "https://www.zhipin.com/job_detail/hangzhou.html",
                    "job_type": "boss_browser",
                },
                {
                    "title": "Python 搜索岗位",
                    "company": "搜索公司",
                    "city": "上海",
                    "salary": "18-25K",
                    "description": "搜索后的真实岗位，需要 Python。",
                    "url": "https://www.zhipin.com/job_detail/search.html",
                    "job_type": "boss_browser",
                }
            ],
            "diagnostics": {"candidate_card_count": 1},
        }

    monkeypatch.setenv("BROWSER_CDP_URL", "127.0.0.1:9222")
    monkeypatch.setattr(browser_job_extractor_service, "urlopen", fake_urlopen)
    monkeypatch.setattr(browser_job_extractor_service, "sleep", lambda seconds: None)
    monkeypatch.setattr(browser_job_extractor_service.CdpRuntimeClient, "evaluate", fake_evaluate)
    app = create_app(db_path=tmp_path / "agent-business.sqlite3")
    client = TestClient(app)

    response = client.post(
        "/api/platform-jobs/search",
        json={"platforms": ["boss"], "keywords": ["Python 实习"], "city": "上海", "limit": 5},
    )

    assert response.status_code == 200
    payload = response.json()
    assert len(expressions) == 2
    assert "Python 实习" in expressions[0]
    assert "上海" in expressions[0]
    assert "101020100" in expressions[0]
    assert "document.querySelectorAll" in expressions[1]
    assert len(payload["extractions"][0]["jobs"]) == 1
    assert payload["extractions"][0]["jobs"][0]["title"] == "Python 搜索岗位"
    assert payload["extractions"][0]["jobs"][0]["salary"] == "18-25K"


def test_platform_job_extraction_flags_polluted_text_and_hides_untrusted_fields(tmp_path, monkeypatch):
    class FakeCdpResponse:
        def __enter__(self):
            return self

        def __exit__(self, exc_type, exc, traceback):
            return False

        def read(self):
            return json.dumps(
                [
                    {
                        "type": "page",
                        "url": "https://www.zhipin.com/web/geek/job?query=Python",
                        "webSocketDebuggerUrl": "ws://127.0.0.1:9222/devtools/page/1",
                    }
                ]
            ).encode("utf-8")

    def fake_evaluate(self, websocket_url, expression):
        return {
            "jobs": [
                {
                    "title": "Python 后端实习生",
                    "company": "真实公司",
                    "city": "上海",
                    "salary": "□□K·□□薪",
                    "description": "真实页面可见岗位内容，需要 Python。",
                    "url": "https://www.zhipin.com/job_detail/abc.html",
                    "job_type": "boss_browser",
                },
                {
                    "title": "□□□",
                    "company": "□□",
                    "city": "□□",
                    "salary": "□□",
                    "description": "□□□",
                    "url": "https://www.zhipin.com/job_detail/bad.html",
                    "job_type": "boss_browser",
                },
                {
                    "title": "□□□□开发□□",
                    "company": "可读公司",
                    "city": "上海",
                    "salary": "薪资面议",
                    "description": "描述可读但标题污染，不应进入岗位池。",
                    "url": "https://www.zhipin.com/job_detail/partial-bad.html",
                    "job_type": "boss_browser",
                },
            ],
            "diagnostics": {
                "matched_selector_counts": {".job-card-wrapper": 2},
                "candidate_card_count": 2,
            },
        }

    monkeypatch.delenv("BROWSER_CDP_URL", raising=False)
    monkeypatch.setattr(browser_job_extractor_service, "urlopen", lambda url, timeout: FakeCdpResponse())
    monkeypatch.setattr(browser_job_extractor_service.CdpRuntimeClient, "evaluate", fake_evaluate)
    app = create_app(db_path=tmp_path / "agent-business.sqlite3")
    client = TestClient(app)

    response = client.post("/api/platform-jobs/extract", json={"platforms": ["boss"], "limit": 5})

    assert response.status_code == 200
    extraction = response.json()["extractions"][0]
    assert extraction["status"] == "success"
    assert len(extraction["jobs"]) == 1
    assert extraction["jobs"][0]["title"] == "Python 后端实习生"
    assert extraction["jobs"][0]["salary"] == "薪资读取失败"
    assert extraction["diagnostics"]["extracted_job_count"] == 1
    assert extraction["diagnostics"]["text_quality_warnings"]
    assert "dropped polluted job" in " ".join(extraction["diagnostics"]["text_quality_warnings"])


def test_platform_job_extraction_drops_placeholder_titles_and_infers_city_salary(tmp_path, monkeypatch):
    class FakeCdpResponse:
        def __enter__(self):
            return self

        def __exit__(self, exc_type, exc, traceback):
            return False

        def read(self):
            return json.dumps(
                [
                    {
                        "type": "page",
                        "url": "https://www.zhipin.com/web/geek/job?query=Python",
                        "webSocketDebuggerUrl": "ws://127.0.0.1:9222/devtools/page/1",
                    }
                ]
            ).encode("utf-8")

    def fake_evaluate(self, websocket_url, expression):
        return {
            "jobs": [
                {
                    "title": "口口口开发口口（金服科技方向）",
                    "company": "未知公司",
                    "city": "上海",
                    "salary": "薪资面议",
                    "description": "标题混入大量占位符，不能作为可信岗位入库。",
                    "url": "https://www.zhipin.com/job_detail/polluted-title.html",
                    "job_type": "boss_browser",
                },
                {
                    "title": "图像算法实习",
                    "company": "可信公司",
                    "city": "□□",
                    "salary": "",
                    "salary_candidates": ["图像算法实习 可信公司 上海 180-200/天 4天/周 2个月"],
                    "description": "岗位要求：负责图像算法实验、数据分析和模型评估，工作地点上海。",
                    "url": "https://www.zhipin.com/job_detail/vision.html",
                    "job_type": "boss_browser",
                },
            ],
            "diagnostics": {
                "matched_selector_counts": {".job-card-wrapper": 2},
                "candidate_card_count": 2,
            },
        }

    monkeypatch.setenv("BROWSER_CDP_URL", "127.0.0.1:9222")
    monkeypatch.setattr(browser_job_extractor_service, "urlopen", lambda url, timeout: FakeCdpResponse())
    monkeypatch.setattr(browser_job_extractor_service.CdpRuntimeClient, "evaluate", fake_evaluate)
    app = create_app(db_path=tmp_path / "agent-business.sqlite3")
    client = TestClient(app)

    response = client.post("/api/platform-jobs/extract", json={"platforms": ["boss"], "limit": 5})

    assert response.status_code == 200
    extraction = response.json()["extractions"][0]
    assert len(extraction["jobs"]) == 1
    job = extraction["jobs"][0]
    assert job["title"] == "图像算法实习"
    assert job["city"] == "上海"
    assert job["salary"] == "180-200/天"
    warnings = " ".join(extraction["diagnostics"]["text_quality_warnings"])
    assert "dropped polluted job" in warnings


def test_shixiseng_extraction_recovers_company_salary_and_detail_from_candidates(tmp_path, monkeypatch):
    class FakeCdpResponse:
        def __enter__(self):
            return self

        def __exit__(self, exc_type, exc, traceback):
            return False

        def read(self):
            return json.dumps(
                [
                    {
                        "type": "page",
                        "url": "https://www.shixiseng.com/interns",
                        "webSocketDebuggerUrl": "ws://127.0.0.1:9222/devtools/page/1",
                    }
                ]
            ).encode("utf-8")

    def fake_evaluate(self, websocket_url, expression):
        assert "company_candidates" in expression
        return {
            "jobs": [
                {
                    "title": "口口后端开发口口",
                    "company": "未知公司",
                    "city": "上海",
                    "salary": "薪资面议",
                    "description": "标题占位符污染，不能进入岗位池。",
                    "url": "https://www.shixiseng.com/intern/bad",
                    "job_type": "shixiseng_browser",
                },
                {
                    "title": "图像算法实习",
                    "company": "未知公司",
                    "company_candidates": [
                        "图像算法实习 光大养老 医疗/健康/制药/生物 本科及以上 北京 180-200/天"
                    ],
                    "city": "□□",
                    "salary": "",
                    "salary_candidates": [
                        "图像算法实习 光大养老 医疗/健康/制药/生物 本科及以上 北京 180-200/天 4天/周 2个月"
                    ],
                    "description": "列表卡片摘要：图像算法实习 光大养老 北京 180-200/天",
                    "detail_description": (
                        "岗位要求：负责图像算法实验、数据清洗、模型评估和结果可视化。"
                        "任职要求：熟悉 Python、SQL 和基础机器学习。"
                    ),
                    "url": "https://www.shixiseng.com/intern/vision",
                    "job_type": "shixiseng_browser",
                    "detail_status": "detail_fetched",
                    "detail_reason": "详情页已补全岗位要求。",
                },
            ],
            "diagnostics": {
                "matched_selector_counts": {".intern-wrap": 2},
                "candidate_card_count": 2,
            },
        }

    monkeypatch.setenv("BROWSER_CDP_URL", "127.0.0.1:9222")
    monkeypatch.setattr(browser_job_extractor_service, "urlopen", lambda url, timeout: FakeCdpResponse())
    monkeypatch.setattr(browser_job_extractor_service.CdpRuntimeClient, "evaluate", fake_evaluate)
    app = create_app(db_path=tmp_path / "agent-business.sqlite3")
    client = TestClient(app)

    response = client.post("/api/platform-jobs/extract", json={"platforms": ["shixiseng"], "limit": 5})

    assert response.status_code == 200
    extraction = response.json()["extractions"][0]
    assert extraction["status"] == "success"
    assert len(extraction["jobs"]) == 1
    job = extraction["jobs"][0]
    assert job["title"] == "图像算法实习"
    assert job["company"] == "光大养老"
    assert job["city"] == "北京"
    assert job["salary"] == "180-200/天"
    assert job["detail_status"] == "detail_fetched"
    assert "任职要求" in job["description"]
    warnings = " ".join(extraction["diagnostics"]["text_quality_warnings"])
    assert "dropped polluted job" in warnings


def test_platform_job_extraction_prefers_valid_salary_candidates_over_polluted_salary(tmp_path, monkeypatch):
    class FakeCdpResponse:
        def __enter__(self):
            return self

        def __exit__(self, exc_type, exc, traceback):
            return False

        def read(self):
            return json.dumps(
                [
                    {
                        "type": "page",
                        "url": "https://www.zhipin.com/web/geek/job?query=Python",
                        "webSocketDebuggerUrl": "ws://127.0.0.1:9222/devtools/page/1",
                    }
                ]
            ).encode("utf-8")

    def fake_evaluate(self, websocket_url, expression):
        assert "salary_candidates" in expression
        return {
            "jobs": [
                {
                    "title": "算法工程师",
                    "company": "真实公司",
                    "city": "杭州",
                    "salary": "□□K·□□薪",
                    "salary_candidates": ["□□K·□□薪", "算法工程师 杭州 15-25K·14薪 经验不限"],
                    "description": "岗位卡片可见内容，需要 Python。",
                    "url": "https://www.zhipin.com/job_detail/real.html",
                    "job_type": "boss_browser",
                },
                {
                    "title": "数据分析实习生",
                    "company": "另一家公司",
                    "city": "上海",
                    "salary": "",
                    "salary_candidates": ["数据分析实习生 上海 经验不限"],
                    "description": "岗位卡片没有展示薪资。",
                    "url": "https://www.zhipin.com/job_detail/no-salary.html",
                    "job_type": "boss_browser",
                },
            ],
            "diagnostics": {
                "matched_selector_counts": {".job-card-wrapper": 2},
                "candidate_card_count": 2,
            },
        }

    monkeypatch.setenv("BROWSER_CDP_URL", "127.0.0.1:9222")
    monkeypatch.setattr(browser_job_extractor_service, "urlopen", lambda url, timeout: FakeCdpResponse())
    monkeypatch.setattr(browser_job_extractor_service.CdpRuntimeClient, "evaluate", fake_evaluate)
    app = create_app(db_path=tmp_path / "agent-business.sqlite3")
    client = TestClient(app)

    response = client.post("/api/platform-jobs/extract", json={"platforms": ["boss"], "limit": 5})

    assert response.status_code == 200
    jobs = response.json()["extractions"][0]["jobs"]
    assert jobs[0]["salary"] == "15-25K·14薪"
    assert jobs[1]["salary"] == "薪资读取失败"
    warnings = response.json()["extractions"][0]["diagnostics"]["text_quality_warnings"]
    assert not any("hid polluted salary" in warning for warning in warnings)


def test_boss_extraction_recovers_live_card_salary_company_and_drops_search_widgets(tmp_path, monkeypatch):
    class FakeCdpResponse:
        def __enter__(self):
            return self

        def __exit__(self, exc_type, exc, traceback):
            return False

        def read(self):
            return json.dumps(
                [
                    {
                        "type": "page",
                        "url": "https://www.zhipin.com/web/geek/jobs?query=Python&city=101020100",
                        "webSocketDebuggerUrl": "ws://127.0.0.1:9222/devtools/page/1",
                    }
                ]
            ).encode("utf-8")

    boss_loading_detail = "BOSS正在加载中... function getCookie(e){} window.Promise = window.Promise || {}"

    def fake_evaluate(self, websocket_url, expression):
        assert ".boss-name" in expression
        assert ".company-location" in expression
        return {
            "jobs": [
                {
                    "title": "职位搜索",
                    "company": "BOSS直聘APP",
                    "city": "工作区域",
                    "salary": "",
                    "description": "首页 职位 公司 校园 海归 APP 有了 海外",
                    "url": "https://www.zhipin.com/job_detail/",
                    "job_type": "boss_browser",
                },
                {
                    "title": "Python",
                    "company": "红烨达信息科技",
                    "company_candidates": ["红烨达信息科技", "Python \ue032\ue036\ue031-\ue033\ue036\ue031元/天 红烨达信息科技 上海·杨浦区·五角场"],
                    "city": "上海·杨浦区·五角场",
                    "salary": "\ue032\ue036\ue031-\ue033\ue036\ue031元/天",
                    "salary_candidates": ["Python \ue032\ue036\ue031-\ue033\ue036\ue031元/天 4天/周 6个月 本科"],
                    "description": "Python \ue032\ue036\ue031-\ue033\ue036\ue031元/天 4天/周 6个月 本科 红烨达信息科技 上海·杨浦区·五角场",
                    "detail_description": boss_loading_detail,
                    "detail_status": "detail_fetched",
                    "detail_reason": "详情页已补全岗位要求。",
                    "url": "https://www.zhipin.com/job_detail/python.html",
                    "job_type": "boss_browser",
                },
                {
                    "title": "agent开发实习生",
                    "company": "矩阵起源",
                    "company_candidates": ["矩阵起源", "agent开发实习生 \ue033\ue03a\ue031-\ue034\ue031\ue031元/天 矩阵起源 上海·杨浦区·五角场"],
                    "city": "上海·杨浦区·五角场",
                    "salary": "\ue033\ue03a\ue031-\ue034\ue031\ue031元/天",
                    "salary_candidates": ["agent开发实习生 \ue033\ue03a\ue031-\ue034\ue031\ue031元/天 4天/周 4个月 本科"],
                    "description": "agent开发实习生 \ue033\ue03a\ue031-\ue034\ue031\ue031元/天 4天/周 4个月 本科 矩阵起源 上海·杨浦区·五角场",
                    "detail_description": boss_loading_detail,
                    "detail_status": "detail_fetched",
                    "detail_reason": "详情页已补全岗位要求。",
                    "url": "https://www.zhipin.com/job_detail/agent.html",
                    "job_type": "boss_browser",
                },
            ],
            "diagnostics": {
                "matched_selector_counts": {".job-card-box": 15},
                "candidate_card_count": 15,
            },
        }

    monkeypatch.setenv("BROWSER_CDP_URL", "127.0.0.1:9222")
    monkeypatch.setattr(browser_job_extractor_service, "urlopen", lambda url, timeout: FakeCdpResponse())
    monkeypatch.setattr(browser_job_extractor_service.CdpRuntimeClient, "evaluate", fake_evaluate)
    app = create_app(db_path=tmp_path / "agent-business.sqlite3")
    client = TestClient(app)

    response = client.post("/api/platform-jobs/extract", json={"platforms": ["boss"], "limit": 5})

    assert response.status_code == 200
    extraction = response.json()["extractions"][0]
    assert extraction["status"] == "success"
    assert extraction["diagnostics"]["candidate_card_count"] == 15
    assert [job["title"] for job in extraction["jobs"]] == ["Python", "agent开发实习生"]
    assert extraction["jobs"][0]["company"] == "红烨达信息科技"
    assert extraction["jobs"][0]["salary"] == "150-250元/天"
    assert extraction["jobs"][0]["detail_status"] == "card_only"
    assert "BOSS正在加载中" not in extraction["jobs"][0]["description"]
    assert extraction["jobs"][1]["salary"] == "290-300元/天"


def test_platform_job_extraction_prefers_detail_page_requirements_and_salary(tmp_path, monkeypatch):
    class FakeCdpResponse:
        def __enter__(self):
            return self

        def __exit__(self, exc_type, exc, traceback):
            return False

        def read(self):
            return json.dumps(
                [
                    {
                        "type": "page",
                        "url": "https://www.zhipin.com/web/geek/job?query=Agent",
                        "webSocketDebuggerUrl": "ws://127.0.0.1:9222/devtools/page/1",
                    }
                ]
            ).encode("utf-8")

    def fake_evaluate(self, websocket_url, expression):
        assert "detail_salary_candidates" in expression
        return {
            "jobs": [
                {
                    "title": "AI Agent优化工程师实习",
                    "company": "真实公司",
                    "city": "上海",
                    "salary": "薪资读取失败",
                    "description": "AI Agent优化工程师实习 - 元/天 4天/周 硕士 上海",
                    "detail_description": (
                        "职位描述 岗位职责：负责 Agent 工具链评测、提示词优化和数据分析。"
                        "任职要求：熟悉 Python、SQL、React，了解大模型调用和实验记录。"
                    ),
                    "detail_salary_candidates": ["250-350元/天", "4天/周"],
                    "url": "https://www.zhipin.com/job_detail/agent.html",
                    "job_type": "boss_browser",
                }
            ],
            "diagnostics": {
                "matched_selector_counts": {".job-card-wrapper": 1},
                "candidate_card_count": 1,
            },
        }

    monkeypatch.setenv("BROWSER_CDP_URL", "127.0.0.1:9222")
    monkeypatch.setattr(browser_job_extractor_service, "urlopen", lambda url, timeout: FakeCdpResponse())
    monkeypatch.setattr(browser_job_extractor_service.CdpRuntimeClient, "evaluate", fake_evaluate)
    app = create_app(db_path=tmp_path / "agent-business.sqlite3")
    client = TestClient(app)

    response = client.post("/api/platform-jobs/extract", json={"platforms": ["boss"], "limit": 5})

    assert response.status_code == 200
    job = response.json()["extractions"][0]["jobs"][0]
    assert job["salary"] == "250-350元/天"
    assert job["detail_status"] == "detail_fetched"
    assert job["detail_reason"] == "详情页已补全岗位要求。"
    assert "岗位职责" in job["description"]
    assert "提示词优化" in job["description"]
    assert "AI Agent优化工程师实习 - 元/天" not in job["description"]


def test_detail_refresh_html_regression_extracts_requirements_and_salary():
    html = """
    <html>
      <body>
        <main class="job-detail-box">
          <h1 class="job-name">AI Agent优化工程师实习</h1>
          <a class="company-name">启明数据实验室</a>
          <span class="job-area">上海·浦东新区</span>
          <span class="job-salary">250-350元/天</span>
          <section class="job-sec-text">
            <h3>职位描述</h3>
            <p>岗位职责：负责 Agent 工具链评测、提示词优化、实验数据分析和可视化看板。</p>
            <p>任职要求：熟悉 Python、SQL、React，了解大模型 API 调用和实验记录。</p>
          </section>
        </main>
      </body>
    </html>
    """
    service = browser_job_extractor_service.BrowserJobExtractorService(cdp_url="http://127.0.0.1:9222")

    job = service._normalize_detail_refresh_result(
        "boss",
        "https://www.zhipin.com/job_detail/detail-regression.html",
        {"html": html},
    )

    assert job.title == "AI Agent优化工程师实习"
    assert job.company == "启明数据实验室"
    assert job.city == "上海·浦东新区"
    assert job.salary == "250-350元/天"
    assert job.detail_status == "detail_fetched"
    assert "岗位职责" in job.description
    assert "提示词优化" in job.description
    assert "任职要求" in job.description


def test_platform_job_extraction_script_contains_resilient_selectors(tmp_path, monkeypatch):
    class FakeCdpResponse:
        def __enter__(self):
            return self

        def __exit__(self, exc_type, exc, traceback):
            return False

        def read(self):
            return json.dumps(
                [
                    {
                        "type": "page",
                        "url": "https://www.zhipin.com/web/geek/job?query=Python",
                        "webSocketDebuggerUrl": "ws://127.0.0.1:9222/devtools/page/1",
                    }
                ]
            ).encode("utf-8")

    def fake_evaluate(self, websocket_url, expression):
        assert websocket_url == "ws://127.0.0.1:9222/devtools/page/1"
        assert "fallback_links" in expression
        assert "[data-jobid]" in expression
        assert "[data-job-id]" in expression
        assert "div[class*='job']" in expression
        assert "firstText" in expression
        assert "scriptSalaryCandidates" in expression
        assert "nextElementSibling" in expression
        assert "fetchDetail" in expression
        assert "await Promise.all" in expression
        return {
            "jobs": [
                {
                    "title": "Backend Intern",
                    "company": "Real Company",
                    "city": "Shanghai",
                    "salary": "200-300/day",
                    "description": "Visible browser job card",
                    "url": "https://www.zhipin.com/job_detail/fallback.html",
                    "job_type": "boss_browser",
                }
            ],
            "diagnostics": {
                "matched_selector_counts": {"fallback_links": 2, "[data-jobid]": 1},
                "candidate_card_count": 1,
            },
        }

    monkeypatch.setenv("BROWSER_CDP_URL", "127.0.0.1:9222")
    monkeypatch.setattr(browser_job_extractor_service, "urlopen", lambda url, timeout: FakeCdpResponse())
    monkeypatch.setattr(browser_job_extractor_service.CdpRuntimeClient, "evaluate", fake_evaluate)
    app = create_app(db_path=tmp_path / "agent-business.sqlite3")
    client = TestClient(app)

    response = client.post("/api/platform-jobs/extract", json={"platforms": ["boss"], "limit": 5})

    assert response.status_code == 200
    extraction = response.json()["extractions"][0]
    assert extraction["status"] == "success"
    assert extraction["diagnostics"]["matched_selector_counts"]["fallback_links"] == 2
    assert extraction["diagnostics"]["matched_selector_counts"]["[data-jobid]"] == 1


def test_boss_browser_script_targets_jobs_page_and_waits_for_rendered_cards():
    service = browser_job_extractor_service.BrowserJobExtractorService(cdp_url="http://127.0.0.1:9222")

    search_url = service._search_url("boss", "Python 实习", "上海")
    extract_script = service._extract_script("boss", 10)

    assert "https://www.zhipin.com/web/geek/jobs?" in search_url
    assert "waitForCandidateCards" in extract_script
    assert "setTimeout(resolve, 250)" in extract_script


def test_boss_browser_script_scrolls_and_recovers_cards_from_job_detail_links():
    service = browser_job_extractor_service.BrowserJobExtractorService(cdp_url="http://127.0.0.1:9222")

    extract_script = service._extract_script("boss", 20)

    assert "scrollForMoreCards" in extract_script
    assert "collectCandidateCards" in extract_script
    assert "bossCardRootFromLink" in extract_script
    assert "[class*='jobCard']" in extract_script
    assert "a[href*='/job_detail/']" in extract_script


def test_boss_browser_script_snapshots_cards_during_scroll_for_virtual_lists():
    service = browser_job_extractor_service.BrowserJobExtractorService(cdp_url="http://127.0.0.1:9222")

    extract_script = service._extract_script("boss", 20)

    assert "cardSnapshots" in extract_script
    assert "snapshotCandidateCard" in extract_script
    assert "rawPayloads" in extract_script
    assert "snapshotKey" in extract_script
    assert "cardSnapshots.slice" in extract_script


def test_cdp_runtime_client_extends_socket_timeout_for_long_running_scripts():
    source = inspect.getsource(browser_job_extractor_service.CdpRuntimeClient.evaluate)

    assert "EVALUATE_TIMEOUT_SECONDS" in source
    assert "sock.settimeout(self.EVALUATE_TIMEOUT_SECONDS)" in source


def test_platform_job_extraction_reports_diagnostics_when_tab_is_missing(tmp_path, monkeypatch):
    class FakeCdpResponse:
        def __enter__(self):
            return self

        def __exit__(self, exc_type, exc, traceback):
            return False

        def read(self):
            return json.dumps(
                [
                    {
                        "type": "page",
                        "url": "https://example.com/",
                    }
                ]
            ).encode("utf-8")

    monkeypatch.setenv("BROWSER_CDP_URL", "127.0.0.1:9222")
    monkeypatch.setattr(browser_job_extractor_service, "urlopen", lambda url, timeout: FakeCdpResponse())
    app = create_app(db_path=tmp_path / "agent-business.sqlite3")
    client = TestClient(app)

    response = client.post("/api/platform-jobs/extract", json={"platforms": ["boss"], "limit": 5})

    assert response.status_code == 200
    extraction = response.json()["extractions"][0]
    assert extraction["status"] == "tab_not_found"
    assert extraction["diagnostics"]["tab_detected"] is False
    assert extraction["diagnostics"]["failure_reason"] == "未检测到平台标签页"
    assert "打开 BOSS" in extraction["diagnostics"]["suggestion"]
