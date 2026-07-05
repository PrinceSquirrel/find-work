from __future__ import annotations

import os
import shutil
import subprocess
import tempfile
from io import BytesIO
from pathlib import Path

from docx import Document
from pypdf import PdfReader


PROJECT_HEADINGS = ("项目经历", "项目经验", "项目实践", "项目")
LOCKED_EDUCATION_PREFIXES = ("教育经历", "教育背景", "教育", "学历")
EDITABLE_SECTION_PREFIXES = (
    "技能",
    "项目",
    "实习",
    "工作",
    "经历",
    "自我",
    "个人",
    "证书",
    "获奖",
    "校园",
)
SECTION_PREFIXES = (
    "教育",
    "技能",
    "实习",
    "工作",
    "校园",
    "证书",
    "获奖",
    "自我",
    "个人",
)


class DocxProjectTemplateService:
    def render(self, template_bytes: bytes, resume_rewrite: str) -> bytes:
        document = Document(BytesIO(template_bytes))
        rewrite_text = self._normalize_rewrite_text(resume_rewrite)
        self._replace_editable_resume_body(document, rewrite_text)
        output = BytesIO()
        document.save(output)
        return output.getvalue()

    def _replace_editable_resume_body(self, document, rewrite_text: str) -> None:
        paragraphs = list(document.paragraphs)
        start_index = self._editable_start_index(paragraphs)
        if start_index is None:
            self._replace_project_section(document, rewrite_text)
            return
        first_replaced = False
        for paragraph in paragraphs[start_index:]:
            if self._paragraph_has_media(paragraph):
                continue
            if not first_replaced:
                paragraph.text = rewrite_text
                first_replaced = True
            else:
                paragraph.text = ""
        if not first_replaced:
            document.add_paragraph(rewrite_text)

    def _editable_start_index(self, paragraphs) -> int | None:
        education_index = None
        for index, paragraph in enumerate(paragraphs):
            text = paragraph.text.strip()
            if text and any(text.startswith(prefix) for prefix in LOCKED_EDUCATION_PREFIXES):
                education_index = index
        if education_index is None:
            return None
        for index in range(education_index + 1, len(paragraphs)):
            paragraph = paragraphs[index]
            text = paragraph.text.strip()
            if self._paragraph_has_media(paragraph):
                continue
            if text and any(text.startswith(prefix) for prefix in EDITABLE_SECTION_PREFIXES):
                return index
        for index in range(education_index + 1, len(paragraphs)):
            if not self._paragraph_has_media(paragraphs[index]):
                return index
        return len(paragraphs)

    def _paragraph_has_media(self, paragraph) -> bool:
        return bool(paragraph._element.xpath(".//w:drawing | .//w:pict"))

    def _replace_project_section(self, document, project_text: str) -> None:
        paragraphs = list(document.paragraphs)
        start_index = self._project_start_index(paragraphs)
        if start_index is None:
            raise ValueError("未找到可编辑简历正文，无法安全套用原简历模板")
        end_index = self._section_end_index(paragraphs, start_index)
        if start_index >= len(paragraphs):
            document.add_paragraph(project_text)
            return
        paragraphs[start_index].text = project_text
        for paragraph in paragraphs[start_index + 1 : end_index]:
            paragraph.text = ""

    def _project_start_index(self, paragraphs) -> int | None:
        for index, paragraph in enumerate(paragraphs):
            text = paragraph.text.strip()
            if not text:
                continue
            if text in PROJECT_HEADINGS:
                return index + 1
            if text.startswith(("项目:", "项目：")):
                return index
            if any(text.startswith(f"{heading}:") or text.startswith(f"{heading}：") for heading in PROJECT_HEADINGS):
                return index
        return None

    def _section_end_index(self, paragraphs, start_index: int) -> int:
        for index in range(start_index + 1, len(paragraphs)):
            text = paragraphs[index].text.strip()
            if text and any(text.startswith(prefix) for prefix in SECTION_PREFIXES):
                return index
        return len(paragraphs)

    def _normalize_rewrite_text(self, resume_rewrite: str) -> str:
        lines = [line.strip() for line in resume_rewrite.splitlines() if line.strip()]
        if not lines:
            raise ValueError("简历改写内容为空，无法生成模板化简历")
        return "\n".join(lines)


class DocxToPdfConverter:
    def convert(self, docx_bytes: bytes) -> bytes:
        with tempfile.TemporaryDirectory(prefix="agent-business-pdf-") as temp_dir:
            temp_path = Path(temp_dir)
            docx_path = temp_path / "tailored-resume.docx"
            pdf_path = temp_path / "tailored-resume.pdf"
            docx_path.write_bytes(docx_bytes)
            if os.name == "nt" and self._convert_with_word(docx_path, pdf_path):
                return pdf_path.read_bytes()
            if self._convert_with_libreoffice(docx_path, temp_path):
                return pdf_path.read_bytes()
        raise RuntimeError("缺少 DOCX 到 PDF 转换器：请安装 Microsoft Word 或 LibreOffice")

    def _convert_with_word(self, docx_path: Path, pdf_path: Path) -> bool:
        word = None
        document = None
        try:
            import win32com.client  # type: ignore

            word = win32com.client.DispatchEx("Word.Application")
            word.Visible = False
            document = word.Documents.Open(str(docx_path))
            document.ExportAsFixedFormat(str(pdf_path), 17)
            return pdf_path.exists()
        except Exception:
            return False
        finally:
            if document is not None:
                document.Close(False)
            if word is not None:
                word.Quit()

    def _convert_with_libreoffice(self, docx_path: Path, output_dir: Path) -> bool:
        executable = shutil.which("soffice") or shutil.which("libreoffice")
        if not executable:
            return False
        try:
            result = subprocess.run(
                [
                    executable,
                    "--headless",
                    "--convert-to",
                    "pdf",
                    "--outdir",
                    str(output_dir),
                    str(docx_path),
                ],
                capture_output=True,
                text=True,
                timeout=60,
            )
        except (OSError, subprocess.TimeoutExpired):
            return False
        return result.returncode == 0 and (output_dir / "tailored-resume.pdf").exists()


class PdfRenderValidator:
    def __init__(self, renderer_executable: str | None = None) -> None:
        self.renderer_executable = renderer_executable

    def validate(self, pdf_bytes: bytes, max_pages: int = 1) -> int:
        try:
            page_count = len(PdfReader(BytesIO(pdf_bytes)).pages)
        except Exception as exc:
            raise RuntimeError(f"invalid PDF output: {type(exc).__name__}") from exc
        if page_count < 1:
            raise RuntimeError("invalid PDF output: no readable pages")
        if page_count > max_pages:
            raise RuntimeError(f"template resume PDF exceeds {max_pages} page")
        self._render_first_page(pdf_bytes)
        return page_count

    def _render_first_page(self, pdf_bytes: bytes) -> None:
        executable = self.renderer_executable or shutil.which("pdftoppm")
        if not executable:
            return
        if not self._renderer_is_available(executable):
            if self.renderer_executable:
                raise RuntimeError("PDF render validation failed: renderer unavailable")
            return
        with tempfile.TemporaryDirectory(prefix="agent-business-pdf-render-") as temp_dir:
            temp_path = Path(temp_dir)
            pdf_path = temp_path / "resume.pdf"
            output_prefix = temp_path / "render"
            pdf_path.write_bytes(pdf_bytes)
            try:
                result = subprocess.run(
                    [
                        executable,
                        "-png",
                        "-f",
                        "1",
                        "-l",
                        "1",
                        str(pdf_path),
                        str(output_prefix),
                    ],
                    capture_output=True,
                    text=True,
                    timeout=30,
                )
            except (OSError, subprocess.TimeoutExpired) as exc:
                raise RuntimeError(f"PDF render validation failed: {type(exc).__name__}") from exc
            rendered_pages = sorted(temp_path.glob("render-*.png"))
            if result.returncode != 0 or not rendered_pages:
                stderr = (result.stderr or result.stdout or "no render output").strip()
                raise RuntimeError(f"PDF render validation failed: {stderr[:180]}")
            if any(page.stat().st_size == 0 for page in rendered_pages):
                raise RuntimeError("PDF render validation failed: empty rendered page")

    def _renderer_is_available(self, executable: str) -> bool:
        try:
            result = subprocess.run(
                [executable, "-v"],
                capture_output=True,
                text=True,
                timeout=10,
            )
        except (OSError, subprocess.TimeoutExpired):
            return False
        return result.returncode == 0


class TailoredResumePdfService:
    def __init__(
        self,
        template_service: DocxProjectTemplateService | None = None,
        converter: DocxToPdfConverter | None = None,
        validator: PdfRenderValidator | None = None,
    ) -> None:
        self.template_service = template_service or DocxProjectTemplateService()
        self.converter = converter or DocxToPdfConverter()
        self.validator = validator or PdfRenderValidator()

    def render(self, tailored_bundle: dict[str, object], template_bytes: bytes) -> bytes:
        resume_rewrite = str(
            tailored_bundle.get("resume_rewrite")
            or tailored_bundle.get("project_rewrite")
            or tailored_bundle.get("resume_text")
            or ""
        )
        docx_bytes = self.template_service.render(template_bytes, resume_rewrite)
        pdf_bytes = self.converter.convert(docx_bytes)
        if self._is_valid_one_page_pdf(pdf_bytes):
            return pdf_bytes
        compressed_rewrite = self._compress_rewrite(resume_rewrite)
        compressed_docx = self.template_service.render(template_bytes, compressed_rewrite)
        compressed_pdf = self.converter.convert(compressed_docx)
        if self._is_valid_one_page_pdf(compressed_pdf):
            return compressed_pdf
        raise RuntimeError("模板化简历超过 1 页，请人工精简可编辑简历正文后重新生成")

    def _is_valid_one_page_pdf(self, pdf_bytes: bytes) -> bool:
        try:
            self.validator.validate(pdf_bytes, max_pages=1)
            return True
        except RuntimeError as exc:
            if "exceeds 1 page" in str(exc):
                return False
            raise

    def _page_count(self, pdf_bytes: bytes) -> int:
        return len(PdfReader(BytesIO(pdf_bytes)).pages)

    def _compress_rewrite(self, resume_rewrite: str) -> str:
        lines = [line.strip() for line in resume_rewrite.splitlines() if line.strip()]
        compact_lines = lines[:8] if lines else [resume_rewrite.strip()]
        return "\n".join(line[:90] for line in compact_lines if line)
